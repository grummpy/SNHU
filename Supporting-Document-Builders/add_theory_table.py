from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parent
SOURCE = BASE_DIR / "input" / "CS 505 Module Five Activity Completed.docx"
OUTPUT = BASE_DIR / "output" / "CS 505 Module Five Activity with Theory Table.docx"
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document(SOURCE)
doc.add_page_break()
heading = doc.add_paragraph("Appendix: Algorithm Theory and Solution Summary", style="Heading 2")
heading.paragraph_format.keep_with_next = True

intro = doc.add_paragraph(
    "The table connects each algorithm’s theoretical basis to the concrete steps used in the examples. "
    "The companion Python and Jupyter files implement the same procedures with plain Python data structures."
)
intro.paragraph_format.space_after = Pt(8)

headers = ["Algorithm", "Theory / invariant", "How the problem is solved", "Implementation and cost"]
rows = [
    (
        "Prim",
        "Cut property: the lightest edge crossing from the current tree to an unvisited vertex is safe. The tree remains connected and acyclic after every choice.",
        "Start at one vertex. Store all crossing edges in a min-heap. Repeatedly accept the lowest-weight edge whose endpoint is new, then add that vertex’s outgoing edges. Stop after V-1 accepted edges.",
        "Adjacency list + binary heap. O(E log V) time and O(V + E) space. A disconnected graph requires restarting once per component."
    ),
    (
        "Kruskal",
        "Cut property plus cycle avoidance: the globally lightest edge joining two different components is safe. Accepted edges always form a forest.",
        "Sort edges by weight. Use disjoint sets to test whether endpoints have different representatives. Accept only edges that merge components. Path compression and union by rank keep the set trees shallow.",
        "Sorted edge list + DSU. O(E log E) time, O(V + E) space, and nearly constant amortized DSU operations. Naturally returns an MSF when disconnected."
    ),
    (
        "Ford-Fulkerson",
        "Residual-capacity invariant: every augmentation preserves capacity limits and flow conservation. Reverse residual edges allow earlier choices to be canceled. No augmenting path implies optimality by max-flow min-cut.",
        "Find an s-to-t path in the residual graph. Send the path bottleneck amount, reduce forward residual capacity, and increase reverse capacity. Repeat until t is unreachable. The provided code uses BFS path selection.",
        "Residual adjacency structure + BFS. Generic integer-capacity bound O(EF); the supplied Edmonds-Karp selection gives O(VE²) time and O(V + E) space."
    ),
]

table = doc.add_table(rows=1, cols=4)
table.autofit = False
table.style = "Normal Table"
widths = [Inches(0.9), Inches(1.75), Inches(2.35), Inches(1.5)]

for i, text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.width = widths[i]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = text
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "003B71")
    cell._tc.get_or_add_tcPr().append(shading)
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for record in rows:
    cells = table.add_row().cells
    for i, text in enumerate(record):
        cells[i].width = widths[i]
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[i].text = text
        for p in cells[i].paragraphs:
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                run.font.size = Pt(8.5)
        if i == 0:
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[i].paragraphs[0].runs[0].bold = True

# Explicit table geometry: 6.5-inch usable width = 9360 DXA.
dxa = [1296, 2520, 3384, 2160]
tbl_pr = table._tbl.tblPr
tbl_borders = OxmlElement("w:tblBorders")
for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "4")
    border.set(qn("w:color"), "A6A6A6")
    tbl_borders.append(border)
tbl_pr.append(tbl_borders)
tbl_w = tbl_pr.find(qn("w:tblW"))
if tbl_w is None:
    tbl_w = OxmlElement("w:tblW")
    tbl_pr.append(tbl_w)
tbl_w.set(qn("w:w"), "9360")
tbl_w.set(qn("w:type"), "dxa")
tbl_ind = OxmlElement("w:tblInd")
tbl_ind.set(qn("w:w"), "0")
tbl_ind.set(qn("w:type"), "dxa")
tbl_pr.append(tbl_ind)
grid = table._tbl.tblGrid
for child in list(grid):
    grid.remove(child)
for width in dxa:
    col = OxmlElement("w:gridCol")
    col.set(qn("w:w"), str(width))
    grid.append(col)
for row in table.rows:
    for i, cell in enumerate(row.cells):
        tcw = cell._tc.get_or_add_tcPr().get_or_add_tcW()
        tcw.set(qn("w:w"), str(dxa[i]))
        tcw.set(qn("w:type"), "dxa")
        margins = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcMar")
        if margins is None:
            margins = OxmlElement("w:tcMar")
            cell._tc.get_or_add_tcPr().append(margins)
        for side in ("top", "left", "bottom", "right"):
            node = OxmlElement(f"w:{side}")
            node.set(qn("w:w"), "90")
            node.set(qn("w:type"), "dxa")
            margins.append(node)

# Repeat the header row if Word moves any table rows to a second page.
tr_pr = table.rows[0]._tr.get_or_add_trPr()
tbl_header = OxmlElement("w:tblHeader")
tbl_header.set(qn("w:val"), "true")
tr_pr.append(tbl_header)

usage = doc.add_paragraph()
usage.paragraph_format.space_before = Pt(10)
usage.paragraph_format.space_after = Pt(4)
run = usage.add_run("Using the companion code")
run.bold = True

for label, text in [
    ("PyCharm: ", "Open graph_algorithms.py and run the file. The demonstration block prints an MST from both algorithms and a maximum flow."),
    ("Jupyter: ", "Keep graph_algorithms.py and graph_algorithms_notebook.ipynb in the same folder, open the notebook, and run cells from top to bottom."),
    ("New inputs: ", "Replace weighted_edges or network. Undirected MST edges use (source, target, weight); flow capacities use nested dictionaries."),
]:
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)

doc.save(OUTPUT)
print(OUTPUT)
