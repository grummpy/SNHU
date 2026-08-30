from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

BASE_DIR = Path(__file__).resolve().parent
SOURCE = BASE_DIR / "input" / "CS 505 Module Five Activity with Theory Table.docx"
OUTPUT = BASE_DIR / "output" / "CS 505 Module Five Activity with Explained Pseudocode.docx"
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document(SOURCE)

prim = """PRIM(G, w, start)                         // Receive the graph, edge weights, and starting vertex.
for each vertex v in V:                    // Initialize information for every vertex.
    key[v] = infinity                      // No connection cost is known yet.
    parent[v] = NIL                        // No MST parent has been chosen yet.
key[start] = 0                             // Make the start vertex the first vertex selected.
Q = min-priority-queue(V, key)             // Store unvisited vertices ordered by cheapest key.
while Q is not empty:                      // Continue until every reachable vertex is processed.
    u = EXTRACT-MIN(Q)                     // Remove the vertex with the cheapest known connection.
    for each edge (u, v) in Adj[u]:        // Examine every neighbor of the selected vertex.
        if v is in Q and w(u,v) < key[v]:  // Improve v only if it is unvisited and this edge is cheaper.
            parent[v] = u                  // Record u as v's MST predecessor.
            key[v] = w(u,v)                // Save the new cheapest connection cost.
            DECREASE-KEY(Q, v, key[v])     // Reorder the priority queue after the improvement.
return {(parent[v], v) : v != start}       // Return the V-1 edges that form the MST."""

kruskal = """KRUSKAL(G, w)                              // Receive the weighted graph.
A = empty set                                // A will hold the selected MST edges.
for each vertex v in V:                      // Prepare one component for every vertex.
    MAKE-SET(v)                              // Initially, each vertex is its own disjoint set.
sort E into nondecreasing order by weight    // Consider the cheapest edges first.
for each edge (u, v) in sorted E:            // Scan every edge in sorted order.
    if FIND-SET(u) != FIND-SET(v):           // Different representatives mean no cycle will form.
        add (u, v) to A                      // Accept this safe, minimum-cost edge.
        UNION(u, v)                          // Merge the two connected components.
return A                                     // Return the MST, or an MSF if G is disconnected."""

dsu = """MAKE-SET(x):                                  // Create a new one-vertex component.
    parent[x] = x                              // Make x its own representative.
    rank[x] = 0                                // A one-node tree begins at rank zero.

FIND-SET(x):                                  // Locate the representative of x's component.
    if parent[x] != x:                         // If x is not the root, follow its parent.
        parent[x] = FIND-SET(parent[x])        // Compress the path by pointing x to the root.
    return parent[x]                           // Return the component representative.

UNION(x, y):                                  // Merge the components containing x and y.
    rx = FIND-SET(x)                           // Find x's representative.
    ry = FIND-SET(y)                           // Find y's representative.
    if rx == ry: return                        // They are already connected; no merge is needed.
    if rank[rx] < rank[ry]:                    // Attach the shorter tree below the taller tree.
        parent[rx] = ry                        // ry becomes the representative.
    else if rank[rx] > rank[ry]:               // Handle the opposite height relationship.
        parent[ry] = rx                        // rx becomes the representative.
    else:                                      // Both trees have the same rank.
        parent[ry] = rx                        // Choose rx as the new representative.
        rank[rx] = rank[rx] + 1                // The combined tree is one level taller.

KRUSKAL-DSU(V, E):                            // Compute an MST/MSF with the disjoint-set structure.
    T = empty set                              // T stores accepted edges.
    for each v in V: MAKE-SET(v)               // Begin with one component per vertex.
    for each (u,v) in E sorted by weight:      // Process edges from cheapest to most expensive.
        if FIND-SET(u) != FIND-SET(v):         // Accept only edges joining different components.
            add (u,v) to T                     // Add the safe edge to the result.
            UNION(u,v)                         // Merge the components after accepting the edge.
    return T                                   // Return the MST or minimum spanning forest."""

def replace_prompt(prefix, intro, blocks):
    paragraph = next(p for p in doc.paragraphs if p.text.startswith(prefix))
    question = paragraph.text.split("\nResponse:")[0]
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    paragraph.add_run(question)
    paragraph.add_run("\nResponse: ").bold = True
    paragraph.add_run(intro)
    for label, code in blocks:
        label_run = paragraph.add_run("\n" + label + "\n")
        label_run.bold = True
        code_run = paragraph.add_run(code)
        code_run.font.name = "Courier New"
        code_run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Courier New")
        code_run.font.size = Pt(8.25)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(8)

replace_prompt(
    "Implement both Prim",
    "Each pseudocode statement is followed by a comment explaining its purpose. Prim grows one tree from a start vertex; Kruskal grows several components and merges them safely.",
    [("Prim's algorithm with line-by-line explanation:", prim),
     ("Kruskal's algorithm with line-by-line explanation:", kruskal)],
)

replace_prompt(
    "Implement Kruskal's algorithm using a disjoint-set",
    "This expanded version shows how path compression and union by rank support Kruskal's cycle test. Each line includes its immediate effect.",
    [("Disjoint-set Kruskal with line-by-line explanation:", dsu)],
)

doc.save(OUTPUT)
print(OUTPUT)
