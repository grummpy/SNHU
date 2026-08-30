from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE_DIR = Path(__file__).resolve().parent
SOURCE = BASE_DIR / "input" / "CS 505 Module Five Activity Graph Theory and Graphical Representations Template.docx"
OUTPUT = BASE_DIR / "output" / "CS 505 Module Five Activity Completed.docx"
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

answers = {
6: {
"text": "Prim’s algorithm (adjacency list and min-priority queue):",
"code": """PRIM(G = (V, E), w, start):
    for each vertex v in V:
        key[v] = infinity; parent[v] = NIL
    key[start] = 0
    Q = min-priority-queue containing all vertices, keyed by key
    while Q is not empty:
        u = EXTRACT-MIN(Q)
        for each edge (u, v) in Adj[u]:
            if v is in Q and w(u, v) < key[v]:
                parent[v] = u
                key[v] = w(u, v)
                DECREASE-KEY(Q, v, key[v])
    return {(parent[v], v) : v != start}""",
"tail": "Kruskal’s algorithm (edge list):",
"code2": """KRUSKAL(G = (V, E), w):
    A = empty set
    for each vertex v in V: MAKE-SET(v)
    sort E into nondecreasing order by weight
    for each edge (u, v) in sorted E:
        if FIND-SET(u) != FIND-SET(v):
            add (u, v) to A
            UNION(u, v)
    return A"""
},
7: {"text": "I traced both algorithms on the following undirected graphs; ties were broken alphabetically. In an unweighted graph, every edge was assigned weight 1.",
"code": """1. Sparse weighted (V=5, E=6)
   Edges: AB1, AC4, BC2, BD5, CD1, DE3
   Prim/Kruskal: AB1, BC2, CD1, DE3; total = 7
2. Dense weighted K5 (V=5, E=10)
   Edges: AB3, AC1, AD6, AE4, BC5, BD2,
          BE7, CD4, CE2, DE5
   Prim/Kruskal: AC1, CE2, AB3, BD2; total = 8
3. Sparse unweighted path P6 (V=6, E=5)
   Prim/Kruskal: all five path edges; total = 5
4. Dense unweighted K5 (V=5, E=10)
   One result: AB, AC, AD, AE; total = 4
5. Disconnected weighted (V=6, E=6)
   Edges: AB1, AC4, BC2; DE1, DF6, EF4
   MSF: AB1, BC2; DE1, EF4; total = 8""",
"tail": "For each connected case, Prim and Kruskal returned the same minimum total weight, although equal-weight graphs allowed different edge sets. On the disconnected case, Kruskal returned the minimum spanning forest. A single run of Prim reached only its starting component; restarting Prim at each unvisited vertex produced the same forest. These tests also illustrate density: the sparse cases have E close to V, while K5 has the maximum 10 undirected edges."
},
8: {"text": "Let V be the number of vertices and E the number of edges. Prim with an adjacency list and binary heap runs in O(E log V) time: each vertex is extracted and edge relaxations may cause heap updates. Its storage is O(V + E) for the graph plus O(V) for keys, parents, membership, and the heap, so total space is O(V + E). With an adjacency matrix and a linear search for the next vertex, Prim runs in O(V²) time and uses O(V²) graph space; this version can be attractive for dense graphs. A Fibonacci-heap formulation has O(E + V log V) amortized time, but a larger implementation constant. Kruskal sorts all edges in O(E log E) time, equivalent to O(E log V) for a simple graph. Its disjoint-set operations add O(E α(V)) time with union by rank and path compression, where α is the extremely slow-growing inverse Ackermann function. Kruskal stores the edge list, output, and disjoint sets in O(E + V) space."
},
9: {"text": "Both algorithms are greedy and produce an MST for a connected, undirected, weighted graph. Prim grows one tree from a selected start and repeatedly chooses the cheapest edge leaving that tree. It fits adjacency-list or matrix representations and is often strong on dense graphs, especially the O(V²) matrix version, because it avoids sorting every edge. Kruskal considers edges globally from smallest to largest and joins components when no cycle is created. It is often strong on sparse graphs and when edges already arrive sorted; its primary cost is sorting. Kruskal naturally produces a forest on disconnected input and can stop after V − 1 accepted edges on connected input. Prim needs an explicit restart to cover multiple components. Actual performance also depends on graph representation, heap implementation, sort quality, memory locality, and tie frequency—not only asymptotic bounds."
},
10: {"text": "MST algorithms minimize the total cost of connecting locations while keeping every location reachable and avoiding redundant cycles. Applications include planning fiber, electrical, water, road, and pipeline backbones; connecting offices or sensors; designing low-cost broadcast trees; clustering data by deleting the largest edges of an MST; approximating solutions to problems such as metric traveling salesperson; and simplifying images or networks while preserving connectivity. The model is appropriate when edge weights capture the relevant cost and one connected backbone is sufficient. If redundancy, direction, capacity, reliability, or multiple objectives matter, an MST is a baseline rather than a complete design because a tree has no backup route after an edge failure."
},
11: {"text": "A disconnected graph has no spanning tree because no edge set can connect vertices that lie in different components. Its minimum spanning forest (MSF) is the union of an MST for each connected component. The forest spans every vertex, contains no cycles, and has minimum total weight among all spanning forests that preserve the original components. If the graph has V vertices and c connected components, an MSF contains V − c edges. Kruskal returns an MSF without modification because it accepts safe edges within components and never invents cross-component edges. Prim can compute the same result by starting a new run from each still-unvisited vertex."
},
12: {"text": "The following version uses both path compression in FIND-SET and union by rank. Path compression makes every visited node point directly to the representative; rank keeps the trees shallow.",
"code": """MAKE-SET(x):
    parent[x] = x; rank[x] = 0

FIND-SET(x):
    if parent[x] != x:
        parent[x] = FIND-SET(parent[x])   // path compression
    return parent[x]

UNION(x, y):
    rx = FIND-SET(x); ry = FIND-SET(y)
    if rx == ry: return
    if rank[rx] < rank[ry]: parent[rx] = ry
    else if rank[rx] > rank[ry]: parent[ry] = rx
    else: parent[ry] = rx; rank[rx] = rank[rx] + 1

KRUSKAL-DSU(V, E):
    T = empty set
    for each v in V: MAKE-SET(v)
    for each (u, v) in E sorted by nondecreasing weight:
        if FIND-SET(u) != FIND-SET(v):
            add (u, v) to T; UNION(u, v)
    return T""",
"tail": "Across m disjoint-set operations, the amortized cost is O(m α(V)); therefore sorting remains Kruskal’s dominant O(E log E) term."
},
16: {"text": "Ford–Fulkerson repeatedly finds an augmenting path in the residual graph. The residual graph includes unused forward capacity and reverse capacity that allows earlier flow choices to be undone.",
"code": """FORD-FULKERSON(G = (V, E), capacity, s, t):
    for each edge (u, v) in E: flow[u, v] = 0
    build residual capacities r[u, v]
    while an s-to-t path P exists in the residual graph:
        delta = minimum r[u, v] over edges (u, v) in P
        for each edge (u, v) in P:
            flow[u, v] = flow[u, v] + delta
            flow[v, u] = flow[v, u] - delta
            r[u, v] = r[u, v] - delta
            r[v, u] = r[v, u] + delta
    return flow and total flow leaving s"""
},
17: {"text": "I traced augmenting paths and updated forward and reverse residual capacities after each augmentation.",
"code": """1. Simple network
   Edges: s→a3, s→b2, a→b1, a→t2, b→t3
   Paths: s-a-t +2; s-b-t +2; s-a-b-t +1
   Maximum flow = 5
2. Two-bottleneck network
   Edges: s→a10, s→b5, a→c4, a→t6, b→c5, c→t8
   Paths: s-a-t +6; s-a-c-t +4; s-b-c-t +4
   Maximum flow = 14
3. Classic six-vertex network
   Edges: s→v1 16, s→v2 13, v1→v2 10, v2→v1 4,
          v1→v3 12, v3→v2 9, v2→v4 14, v4→v3 7,
          v3→t 20, v4→t 4
   Paths: s-v1-v3-t +12; s-v2-v4-t +4;
          s-v2-v4-v3-t +7
   Maximum flow = 23""",
"tail": "In each case, the reported flow obeys capacity constraints and flow conservation. When no augmenting path remained, the vertices reachable from s in the residual graph defined a cut whose capacity equaled the reported flow, independently confirming optimality."
},
18: {"text": "With integer capacities and a path search that scans the graph in O(E), generic Ford–Fulkerson runs in O(EF) time, where F is the value of the maximum flow: each augmentation increases total flow by at least one. This is pseudo-polynomial because F depends on numeric capacity values, not merely the number of input bits. With irrational capacities, arbitrary path selection can even fail to terminate. Choosing breadth-first search produces the Edmonds–Karp variant, which guarantees O(VE²) time regardless of capacity magnitudes. An adjacency-list residual graph uses O(V + E) space; an adjacency matrix uses O(V²). Flow and residual-capacity records require O(E) additional storage, and a path search uses O(V) queue/stack and parent storage."
},
19: {"text": "Ford–Fulkerson is conceptually simple, flexible, and often fast when capacities and the maximum flow are small. Its practical speed is highly sensitive to augmenting-path choice. Poor choices can add only a small amount of flow per iteration or repeatedly cancel earlier choices through reverse edges. Large capacities therefore make the O(EF) bound unattractive. Edmonds–Karp trades some practical overhead for a polynomial guarantee, while Dinic’s algorithm is generally preferred for larger networks. Correct implementations must maintain reverse residual edges, enforce nonnegative residual capacity, preserve flow conservation, and use sufficiently wide numeric types."
},
20: {"text": "Network flow models allocate a limited resource through a capacitated system. Examples include routing data through communication links, distributing water or electricity, shipping goods through warehouses, scheduling jobs on machines, assigning workers or students through bipartite matching, selecting disjoint evacuation routes, allocating airline or rail capacity, and segmenting images with source/sink cuts. The model can also represent baseball elimination and project selection after a suitable graph transformation. Real deployments may require extensions such as minimum-cost flow, multiple sources/sinks (handled with supernodes), lower bounds, time-expanded networks, or fairness and reliability constraints."
},
21: {"text": "An s–t cut partitions the vertices into S and T with s in S and t in T. Its capacity is the sum of capacities of directed edges from S to T. Every feasible s–t flow is at most the capacity of every s–t cut, because all net flow must cross the partition. The max-flow min-cut theorem states that the maximum feasible flow value equals the minimum cut capacity. At Ford–Fulkerson termination, let S be the vertices reachable from s by positive residual-capacity edges and T be the rest. Since t is unreachable, this is a cut; all S-to-T edges are saturated and all T-to-S flow is zero, so the current flow equals the cut capacity. Thus the flow and cut certify each other’s optimality."
},
}

def add_text_run(p, text, bold=False, italic=False, font=None, size=None):
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    if font:
        r.font.name = font
        r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if size: r.font.size = Pt(size)
    return r

doc = Document(SOURCE)
for idx, data in answers.items():
    p = doc.paragraphs[idx]
    question = p.text.split("\n[Insert text.]")[0]
    for run in p.runs:
        run._element.getparent().remove(run._element)
    add_text_run(p, question)
    add_text_run(p, "\nResponse: ", bold=True)
    add_text_run(p, data["text"])
    if data.get("code"):
        add_text_run(p, "\n")
        add_text_run(p, data["code"], font="Courier New", size=9)
    if data.get("tail"):
        add_text_run(p, "\n")
        add_text_run(p, data["tail"])
    if data.get("code2"):
        add_text_run(p, "\n")
        add_text_run(p, data["code2"], font="Courier New", size=9)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.08

# Keep headings with the paragraph that follows and prevent isolated prompt questions.
for i in (4, 14):
    doc.paragraphs[i].paragraph_format.keep_with_next = True
for p in doc.paragraphs:
    p.paragraph_format.widow_control = True

# Ask Word to refresh fields if the template contains any cached fields.
settings = doc.settings._element
update = settings.find(qn('w:updateFields'))
if update is None:
    update = OxmlElement('w:updateFields')
    settings.append(update)
update.set(qn('w:val'), 'true')

doc.core_properties.title = "CS 505 Module Five Activity: Graph Theory and Graphical Representations"
doc.save(OUTPUT)
print(OUTPUT)
