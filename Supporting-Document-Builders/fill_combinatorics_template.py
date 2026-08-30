from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

BASE_DIR = Path(__file__).resolve().parent
SOURCE = BASE_DIR / "input" / "CS 505 Module Five Activity Combinatorics Template.docx"
OUTPUT = BASE_DIR / "output" / "CS 505 Module Five Activity Combinatorics Completed.docx"
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

answers = {
3: "This is a permutation because changing the order creates a different arrangement. There are 10 choices for the first position, 9 for the second, and so on. Therefore, 10! = 10 x 9 x ... x 1 = 3,628,800 arrangements. Jupyter verification: factorial(10).",
4: "A committee is an unordered selection, so use a combination rather than a permutation. C(20,5) = 20!/[5!(20-5)!] = 15,504 possible committees. Jupyter verification: comb(20, 5).",
5: "The order in which the two cards are drawn does not create a different two-card selection, so use C(52,2) = 52!/[2!50!] = (52 x 51)/2 = 1,326 selections. Jupyter verification: comb(52, 2).",
6: "The available multiplicities are M=1, I=4, S=4, and P=2. For every valid count pattern (m,i,s,p) satisfying m+i+s+p=6 and the availability bounds, there are 6!/(m!i!s!p!) distinct arrangements. Summing this expression over all 25 valid patterns gives 1,610 different 6-letter strings. The notebook enumerates the patterns so the bounds are applied correctly.",
7: "Permutations apply when order or assigned positions matter, such as ranking race finishers, arranging a schedule, assigning officers to distinct roles, seating guests, or forming codes without repeated symbols. Combinations apply when only membership matters, such as selecting committees, choosing sampled records, selecting product features, forming teams without roles, or calculating lottery possibilities.",
8: "A permutation counts ordered arrangements; exchanging two selected objects changes the result. Its formula is P(n,r) = n!/(n-r)!. A combination counts unordered selections; the same selected objects in another order are not new. Its formula is C(n,r) = n!/[r!(n-r)!]. For example, selecting Alice then Bob is different from Bob then Alice in a speaking order, but it is the same two-person committee.",
9: "Let xi be the number of balls in box i. Because the balls are identical, the boxes are distinct, and empty boxes are allowed, count the nonnegative solutions to x1+x2+x3+x4+x5+x6=6. Stars and bars gives C(6+6-1,6-1) = C(11,5) = 462 distributions.",
10: "First give each child the required 3 candies, using 9 candies. Let yi be the number of additional candies child i receives. Then y1+y2+y3=6 with yi >= 0. Stars and bars gives C(6+3-1,3-1) = C(8,2) = 28 distributions.",
11: "For k=3 nonnegative variables totaling n=20, stars and bars gives C(n+k-1,k-1) = C(22,2) = 231 solutions. The 20 stars represent units and the two bars divide them among x1, x2, and x3.",
12: "Positive means each variable is at least 1. Set yi=xi-1, so y1+y2+y3=7 with yi >= 0. Stars and bars gives C(7+3-1,3-1) = C(9,2) = 36 positive-integer solutions.",
13: "For lower bounds xi >= ai, assign each minimum first and substitute yi=xi-ai; then apply the ordinary nonnegative formula to the remaining total. Strict positivity is the special case ai=1. Upper bounds can be handled with inclusion-exclusion by subtracting solutions that exceed a limit. Fixed values can be substituted directly, while several irregular restrictions may be handled by conditioning on cases, generating functions, or dynamic programming.",
14: "Stars and bars models allocations of identical units among distinct recipients. Applications include distributing whole-dollar budget units among projects, identical inventory among stores, identical computing tasks among servers, reward tokens among employees, seats among categories, or total event counts among time periods. The method applies when only each recipient's count matters; if items are distinct or capacities are restrictive, the model must be adjusted.",
}

doc = Document(SOURCE)
for idx, answer in answers.items():
    p = doc.paragraphs[idx]
    question = p.text.split("\n[Insert text.]")[0]
    for run in list(p.runs):
        p._p.remove(run._r)
    p.add_run(question)
    p.add_run("\nResponse: ").bold = True
    p.add_run(answer)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.08

doc.core_properties.title = "CS 505 Module Five Activity: Combinatorics"
doc.save(OUTPUT)
print(OUTPUT)
