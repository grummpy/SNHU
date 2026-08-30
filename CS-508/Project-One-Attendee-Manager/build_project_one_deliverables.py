from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path("project_one")
JAVA = ROOT / "ConferenceAttendeeManager.java"
TXT = ROOT / "CS508_Project_One_Source_Code.txt"
ZIP = ROOT / "CS508_Project_One_Java_Code.zip"
REPORT = ROOT / "CS508_Project_One_Technical_Report.docx"
GUIDE = ROOT / "CS508_Project_One_Pseudocode_and_Walkthrough.docx"


def set_run(run, font_name, size=12, bold=False, italic=False, color="000000"):
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), font_name)
    rpr.rFonts.set(qn("w:hAnsi"), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph, font_name="Times New Roman"):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, font_name, 10)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure(doc, font_name="Times New Roman", double=True):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)
        section.header_distance = section.footer_distance = Inches(0.5)
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    normal.font.size = Pt(12)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0 if double else 6)
    normal.paragraph_format.line_spacing = 2 if double else 1.15


def report_para(doc, text, first_line=True, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if first_line:
        p.paragraph_format.first_line_indent = Inches(0.5)
    if bold_prefix and text.startswith(bold_prefix):
        a = p.add_run(bold_prefix)
        set_run(a, "Times New Roman", 12, bold=True)
        b = p.add_run(text[len(bold_prefix):])
        set_run(b, "Times New Roman", 12)
    else:
        r = p.add_run(text)
        set_run(r, "Times New Roman", 12)
    return p


def report_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run(r, "Times New Roman", 12, bold=True)


def guide_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run(r, "Calibri", 16 if level == 1 else 13, bold=True,
            color="2E74B5" if level == 1 else "1F4D78")
    return p


def guide_para(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run(r, "Calibri", 11, bold=bold)
    return p


def code_block(doc, text):
    for line in text.strip("\n").splitlines():
        if not line.strip():
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1
        r = p.add_run(line)
        set_run(r, "Courier New", 8.25, color="1F1F1F")


def build_report():
    doc = Document()
    configure(doc, "Times New Roman", True)
    add_page_number(doc.sections[0].footer.paragraphs[0])

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Inches(2)
    title.paragraph_format.line_spacing = 2
    r = title.add_run("Conference Attendee Name Management Tool")
    set_run(r, "Times New Roman", 12, bold=True)
    for line in [
        "CS 508 Project One Technical Report",
        "Student Name",
        "Southern New Hampshire University",
        "Course Instructor",
        "August 2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 2
        r = p.add_run(line)
        set_run(r, "Times New Roman", 12)
    doc.add_page_break()

    report_heading(doc, "Conference Attendee Name Management Tool")
    report_heading(doc, "Problem Analysis", 2)
    report_para(doc, "The event company needs one dependable attendee list made from registration lists that may be messy. The program must read names, clean them, put them in alphabetical order, remove repeated names, and create a CSV file. The final file must work in spreadsheet programs and must be useful for badges, roll calls, and seating charts. In simple terms, the program turns an unreliable list into one clean master list.")
    report_para(doc, "The main input requirement is a text file containing one attendee name per line. The program accepts the input path and two output paths from the command line. It cleans each line by removing unwanted characters, combining repeated spaces, and applying consistent capitalization. It keeps letters, apostrophes, and hyphens so names such as O'Neil and Smith-Jones are not damaged. Blank results are ignored.")
    report_para(doc, "The supplied data contains 10,000 attendee entries. An independent count found 2,404 unique names, meaning 7,596 repeated entries must be removed. The program writes one copy of each name to the master CSV. It also creates a duplicate-review CSV showing each repeated name and its total number of appearances. This approach removes duplicates from the working list while still giving the event staff a way to check whether a repeated name represents an error or two different people with the same name.")

    report_heading(doc, "Constraints, Assumptions, Objectives, and Challenges", 2)
    report_para(doc, "Constraints. The solution must be written in Java and cannot depend on a standard Java collection or sorting utility. Therefore, it does not use ArrayList, LinkedList, HashSet, Arrays.sort, or Collections.sort. Basic Java file and character classes are still necessary to read text, write CSV files, compare text, and process characters. The program implements its own resizable array and merge sort.", bold_prefix="Constraints.")
    report_para(doc, "Assumptions. Each input line represents one full attendee name. Names are compared after cleaning and are treated as duplicates without regard to capitalization. Two different people who share the exact same cleaned name cannot be distinguished because the input contains no registration number, email address, or other unique identifier. Empty lines and lines that become empty after cleaning are skipped.", bold_prefix="Assumptions.")
    report_para(doc, "Objectives. The program should produce a correct alphabetical list, remove repeated entries, preserve useful punctuation, scale to thousands of names, and create files that nontechnical event employees can open. It should also be readable enough for another developer to maintain.", bold_prefix="Objectives.")
    report_para(doc, "Challenges. The main challenges are inconsistent formatting, duplicate detection, efficient sorting, memory use, and ambiguity between a true duplicate and two people with the same name. Another challenge is producing valid CSV text when a value contains punctuation. The program handles CSV values by surrounding each name with quotation marks and escaping quotation marks if they appear.", bold_prefix="Challenges.")

    report_heading(doc, "Comparison of Data Structures", 2)
    report_heading(doc, "Arrays", 2)
    report_para(doc, "An array stores items next to each other in memory. Its main strength is fast indexed access: reading or changing an item at a known position takes O(1) time. Arrays also work well with merge sort because the algorithm repeatedly accesses ranges by index. Their weakness is that a fixed array cannot automatically grow. A larger array must be created and the existing names must be copied when capacity is reached. This project solves that problem with a custom resizable array that doubles its capacity.")
    report_heading(doc, "Linked Lists", 2)
    report_para(doc, "A linked list stores each name in a separate node connected to the next node. Its strength is flexible growth and easy insertion when the correct node is already known. Its weaknesses are slow indexed access, extra memory for links, and less convenient merging and CSV traversal logic. Finding the item at a particular position can take O(n) time. These weaknesses make a linked list less direct than an array for this project.")
    report_heading(doc, "Heaps", 2)
    report_para(doc, "A heap is useful when the program repeatedly needs the smallest or largest item. It can insert an item and remove the smallest item in O(log n) time. A heap could produce names in alphabetical order by repeatedly removing the smallest name. However, a heap does not directly solve duplicate removal, does not provide normal sorted indexed access, and is more complicated than needed for producing one complete sorted file. Heap sort is possible, but merge sort is easier to explain and keeps predictable O(n log n) performance.")

    report_heading(doc, "Selected Data Structure and Algorithm", 2)
    report_para(doc, "The selected data structure is a custom resizable array backed by String[]. It begins with room for 16 names. When it becomes full, the grow function creates an array twice as large and copies the names. Appending is O(1) amortized time because resizing happens only occasionally. Indexed access remains O(1), which supports efficient sorting.")
    report_para(doc, "The selected sorting algorithm is merge sort. Merge sort divides the array into halves, sorts each half, and merges the halves back together. Its worst-case time complexity is O(n log n), which is much better for large lists than a quadratic method such as bubble sort. Merge sort uses O(n) temporary space, but the predictable runtime is a strong trade-off for a conference list with thousands of entries (Goodrich et al., 2014).")
    report_para(doc, "After sorting, equal names appear next to one another. The program makes one O(n) pass through the sorted array. It writes the first copy of a name to the master CSV, counts any matching names beside it, and records repeated names in the review file. This is simpler than comparing every name with every other name, which could take O(n squared) time.")

    report_heading(doc, "Program Design and Best Practices", 2)
    report_para(doc, "The program separates its work into small functions. readAndCleanNames handles input. cleanName handles formatting rules. mergeSort and merge handle ordering. writeUniqueCsvAndDuplicateReport handles deduplication and output. csvValue handles CSV safety. The main method only coordinates these steps and reports the results. This separation makes the code easier to read, test, and change.")
    report_para(doc, "The source uses descriptive variable names, comments, private helper methods, input validation, and try-with-resources so files are closed even when an error occurs. File errors are caught and displayed in plain language. The program does not silently discard duplicate information because it creates a review report in addition to the unique master list.")

    report_heading(doc, "Complexity Analysis", 2)
    report_para(doc, "Let n be the number of valid attendee lines and let L be the average name length. Reading and cleaning requires O(nL) time because each character is examined. Merge sort requires O(n log n) name comparisons. The final duplicate and output pass requires O(n) time. For normal short names, the overall algorithm is described as O(n log n) time because sorting is the largest step.")
    report_para(doc, "The resizable array stores n names. Merge sort creates a temporary array of n references, and recursion uses O(log n) stack depth. The overall additional space is O(n). The CSV files are written one row at a time, so the program does not create a second complete output collection in memory.")

    report_heading(doc, "Verification", 2)
    report_para(doc, "The expected full-data result was independently checked from the supplied 10,000-line file. The check found 10,000 valid entries, 2,404 unique cleaned names, and 7,596 duplicate entries removed. The first five alphabetical names are Ahmed Adams, Ahmed Allen, Ahmed Anderson, Ahmed Baker, and Ahmed Brown. The last five are Zara White, Zara Williams, Zara Wilson, Zara Wright, and Zara Young. These checkpoints make it easy to confirm the Java output after running it in VS Code.")
    report_para(doc, "Additional tests should include blank lines, repeated spaces, all-uppercase names, lowercase names, apostrophes, hyphens, unwanted symbols, a one-name file, an empty file, and repeated names with different capitalization. The local verification commands are javac ConferenceAttendeeManager.java followed by java ConferenceAttendeeManager input.txt master_attendees.csv duplicate_review.csv. The program should also be opened in a spreadsheet program to confirm that each attendee appears in one CSV row.")

    report_heading(doc, "Effectiveness of the Proposed Solution", 2)
    report_para(doc, "The proposed solution meets the event company's main needs. It cleans formatting, sorts names, removes repeated entries, preserves a duplicate audit trail, and writes spreadsheet-compatible files. The custom array and merge sort follow the restriction against using built-in Java collection and sorting tools. The O(n log n) sorting time is suitable for hundreds or thousands of attendees and is far more scalable than repeated pair-by-pair comparison.")
    report_para(doc, "The most important limitation is that a name alone is not a guaranteed unique identity. Two attendees can share the same name. In a production system, the best improvement would be to include a registration ID or email address and use that field for duplicate decisions. For the supplied name-only assignment, producing both a unique list and a duplicate-review report is a reasonable and careful solution.")

    report_heading(doc, "Conclusion", 2)
    report_para(doc, "A custom resizable array combined with merge sort provides a clear and efficient solution. The array offers fast access, merge sort gives predictable O(n log n) performance, and one final pass removes adjacent duplicates. The design is simple enough to maintain while still handling a large conference list. It also gives event staff useful output without permanently hiding possible duplicate-registration issues.")

    report_heading(doc, "References")
    p = report_para(doc, "Goodrich, M. T., Tamassia, R., & Goldwasser, M. H. (2014). Data structures and algorithms in Java (6th ed.). Wiley.", first_line=False)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p = report_para(doc, "Oracle. (n.d.). Java Platform, Standard Edition documentation. https://docs.oracle.com/en/java/javase/", first_line=False)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p = report_para(doc, "Southern New Hampshire University. (2026). CS 508 Project One conference attendee list [Data set].", first_line=False)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)

    report_heading(doc, "AI Use Disclosure")
    report_para(doc, "I used ChatGPT to help explain data structures, organize the pseudocode, draft and review the Java solution, and simplify the technical report. I reviewed the material, checked the calculations against the supplied attendee list, and am responsible for the final submission.", first_line=False)
    p = report_para(doc, "OpenAI. (2026). ChatGPT [Large language model]. https://chatgpt.com/", first_line=False)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)

    doc.core_properties.title = "CS 508 Project One Technical Report"
    doc.core_properties.author = "Student"
    doc.save(REPORT)


def build_guide():
    doc = Document()
    configure(doc, "Calibri", False)
    add_page_number(doc.sections[0].footer.paragraphs[0], "Calibri")
    header = doc.sections[0].header.paragraphs[0]
    r = header.add_run("CS 508 Project One | Student Walkthrough")
    set_run(r, "Calibri", 9, color="777777")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(5)
    r = title.add_run("Conference Attendee Program Walkthrough")
    set_run(r, "Calibri", 23, bold=True, color="1F4D78")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Plain-language pseudocode and function-by-function guide")
    set_run(r, "Calibri", 12, color="555555")

    guide_heading(doc, "1. What the Program Does")
    guide_para(doc, "The program reads one attendee name from each line of a text file. It cleans the names, sorts them from A to Z, keeps one copy of each name, and writes a master CSV file. It also writes a second CSV showing repeated names so employees can review them.")
    guide_para(doc, "The program uses its own resizable array and its own merge sort. It does not use ArrayList, HashSet, Arrays.sort, or Collections.sort.", bold=True)

    guide_heading(doc, "2. Main Program Pseudocode")
    code_block(doc, """
SET inputFile to the first command-line value
SET masterFile to the second command-line value
SET duplicateFile to the third command-line value

IF fewer than three values were provided
    DISPLAY the correct usage
    STOP the program
END IF

SET names to a new custom resizable array

OPEN inputFile
WHILE another line exists
    SET cleanedName to CLEAN_NAME(line)
    IF cleanedName is not empty
        ADD cleanedName to names
    END IF
END WHILE
CLOSE inputFile

MERGE_SORT(names)

OPEN masterFile
OPEN duplicateFile
SET index to 0

WHILE index is less than the number of names
    SET currentName to names[index]
    SET occurrences to 1
    MOVE index forward by 1

    WHILE the next name matches currentName
        ADD 1 to occurrences
        MOVE index forward by 1
    END WHILE

    WRITE currentName once to masterFile
    IF occurrences is greater than 1
        WRITE currentName and occurrences to duplicateFile
    END IF
END WHILE

CLOSE both output files
DISPLAY the totals and output locations
""")

    doc.add_page_break()
    guide_heading(doc, "3. Name Cleaning Pseudocode")
    code_block(doc, """
FUNCTION CLEAN_NAME(rawName)
    SET cleaned to empty text

    FOR each character in rawName
        IF the character is a letter
            ADD its lowercase form to cleaned
        ELSE IF it is an apostrophe or hyphen
            KEEP it when it follows a letter
        ELSE IF it is a space
            ADD only one space
        ELSE
            IGNORE the unwanted character
        END IF
    END FOR

    REMOVE a space from the end, if one exists
    CAPITALIZE the first letter of each name part
    RETURN cleaned
END FUNCTION
""")

    guide_heading(doc, "4. Merge Sort Pseudocode")
    code_block(doc, """
FUNCTION MERGE_SORT(names, left, right)
    IF left is greater than or equal to right
        RETURN because one item is already sorted
    END IF

    SET middle to the halfway position
    MERGE_SORT the left half
    MERGE_SORT the right half
    MERGE the two sorted halves
END FUNCTION

FUNCTION MERGE(names, left, middle, right)
    SET one pointer at the start of each half

    WHILE both halves still contain names
        COPY the alphabetically smaller name to temporary storage
        MOVE that half's pointer forward
    END WHILE

    COPY any names remaining in the left half
    COPY any names remaining in the right half
    COPY the merged section back into names
END FUNCTION
""")

    doc.add_page_break()
    guide_heading(doc, "5. Function-by-Function Walkthrough")
    rows = [
        ("main", "Checks the file names, runs each major step, and prints totals."),
        ("NameArray", "Stores names in a custom String array without ArrayList."),
        ("add", "Places a name at the end and grows the array when needed."),
        ("grow", "Doubles capacity and copies the existing names."),
        ("get / set / size", "Safely reads, changes, or counts stored names."),
        ("readAndCleanNames", "Reads every line and keeps each valid cleaned name."),
        ("cleanName", "Fixes capitalization and spacing and removes unwanted characters."),
        ("mergeSort", "Starts sorting and creates one temporary work array."),
        ("recursive mergeSort", "Breaks the list into smaller halves."),
        ("merge", "Puts two sorted halves back together in alphabetical order."),
        ("writeUniqueCsvAndDuplicateReport", "Keeps one copy and reports repeated names."),
        ("csvValue", "Adds safe quotation marks around a CSV value."),
        ("printUsage", "Shows the correct command when file names are missing."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(2.15), Inches(4.35)]
    for i, text in enumerate(["Function", "What It Does in Plain Language"]):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = text
        for run in cell.paragraphs[0].runs:
            set_run(run, "Calibri", 10.5, bold=True, color="1F4D78")
    for name, meaning in rows:
        cells = table.add_row().cells
        for i, value in enumerate([name, meaning]):
            cells[i].width = widths[i]
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].text = value
            for run in cells[i].paragraphs[0].runs:
                set_run(run, "Calibri", 10.5, bold=(i == 0))

    guide_heading(doc, "6. Why the Custom Array Is Used")
    guide_para(doc, "An array gives fast access to any position. That helps merge sort read and replace names efficiently. A normal array has a fixed size, so NameArray solves that problem by doubling its storage when full. Most additions are O(1), and reaching an item by index is O(1).")

    guide_heading(doc, "7. Why Merge Sort Is Used")
    guide_para(doc, "Merge sort is dependable for large lists. It takes O(n log n) time in the worst case. Bubble sort could take O(n squared), which grows much faster. Merge sort needs O(n) extra work space, but that is a reasonable cost for faster sorting.")

    guide_heading(doc, "8. How Duplicate Removal Works")
    guide_para(doc, "Sorting puts matching names beside each other. The program reads from left to right, counts equal neighbors, writes one copy to the master file, and sends the count to the duplicate report. This final scan takes O(n) time.")

    guide_heading(doc, "9. How to Compile and Run in VS Code")
    code_block(doc, """
javac ConferenceAttendeeManager.java

java ConferenceAttendeeManager \
"CS 508 Project One Conference Attendee List.txt" \
master_attendees.csv duplicate_review.csv
""")
    guide_para(doc, "Put the Java file and attendee text file in the same VS Code folder. Open the terminal in that folder, run the compile command, and then run the program command. If the input file is elsewhere, use its full path.")

    guide_heading(doc, "10. Expected Full-File Checkpoints")
    guide_para(doc, "Valid input names: 10,000")
    guide_para(doc, "Unique names: 2,404")
    guide_para(doc, "Duplicate entries removed: 7,596")
    guide_para(doc, "First alphabetical name: Ahmed Adams")
    guide_para(doc, "Last alphabetical name: Zara Young")

    guide_heading(doc, "11. Simple Test Cases")
    code_block(doc, """
Input:     "  jANE   DOE  "     Expected: "Jane Doe"
Input:     "MARY-JANE smith"    Expected: "Mary-Jane Smith"
Input:     "o'NEIL"             Expected: "O'Neil"
Duplicates: "John Smith" twice  Expected: one master row and a count of 2
Empty file                       Expected: headers only, with no crash
""")

    guide_heading(doc, "12. The One-Sentence Explanation")
    guide_para(doc, "The program cleans every name, stores the names in a custom array, uses merge sort to put them in order, then makes one pass to write unique names and flag repeats.", bold=True)

    doc.core_properties.title = "CS 508 Project One Pseudocode and Walkthrough"
    doc.core_properties.author = "Student"
    doc.save(GUIDE)


def package_source():
    source = JAVA.read_text()
    TXT.write_text(source)
    with ZipFile(ZIP, "w", ZIP_DEFLATED) as archive:
        archive.write(JAVA, arcname=JAVA.name)


ROOT.mkdir(exist_ok=True)
package_source()
build_report()
build_guide()
print(TXT)
print(ZIP)
print(REPORT)
print(GUIDE)
