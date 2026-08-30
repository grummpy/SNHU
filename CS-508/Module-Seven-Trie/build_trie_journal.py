from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT_DIR = Path(__file__).resolve().parent / 'output'
OUT_DIR.mkdir(parents=True, exist_ok=True)
PNG_PATH = OUT_DIR / 'CS 508 Module Seven Trie Diagram.png'
DOCX_PATH = OUT_DIR / 'CS 508 Module Seven Trie Journal.docx'

WORDS = [
    'lark', 'car', 'lamb', 'carb', 'lad', 'card', 'can', 'ladle',
    'care', 'cared', 'label', 'coma', 'cares', 'lard', 'carat',
    'laugh', 'carp', 'cart', 'luck', 'carry', 'lamp', 'come',
    'cane', 'cone', 'calf'
]

BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
MUTED = '5D6670'
INK = '#1F2933'
GREEN = '#DDF4E4'
GREEN_EDGE = '#247A45'


class TrieNode:
    def __init__(self, prefix='', char='ROOT'):
        self.prefix = prefix
        self.char = char
        self.children = {}
        self.terminal = False
        self.x = 0
        self.y = 0


def build_trie(words):
    root = TrieNode()
    for word in sorted(words):
        node = root
        prefix = ''
        for char in word:
            prefix += char
            if char not in node.children:
                node.children[char] = TrieNode(prefix, char)
            node = node.children[char]
        node.terminal = True
    return root


def count_nodes(node):
    return 1 + sum(count_nodes(child) for child in node.children.values())


def assign_positions(root, width, top, level_gap, side_margin):
    leaves = []

    def collect(node):
        if not node.children:
            leaves.append(node)
        for child in node.children.values():
            collect(child)

    collect(root)
    usable = width - 2 * side_margin
    for index, leaf in enumerate(leaves):
        if len(leaves) == 1:
            leaf.x = width / 2
        else:
            leaf.x = side_margin + index * usable / (len(leaves) - 1)

    def place(node, depth):
        node.y = top + depth * level_gap
        if node.children:
            for child in node.children.values():
                place(child, depth + 1)
            node.x = sum(child.x for child in node.children.values()) / len(node.children)

    place(root, 0)


def create_diagram(root):
    width, height = 3200, 1750
    image = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(image)
    font_dir = Path('/System/Library/Fonts/Supplemental')
    regular = ImageFont.truetype(str(font_dir / 'Arial.ttf'), 34)
    bold = ImageFont.truetype(str(font_dir / 'Arial Bold.ttf'), 35)
    title_font = ImageFont.truetype(str(font_dir / 'Arial Bold.ttf'), 54)
    note_font = ImageFont.truetype(str(font_dir / 'Arial.ttf'), 28)

    assign_positions(root, width, top=250, level_gap=235, side_margin=110)

    title = 'Trie for the 25 Given Words'
    title_box = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((width - (title_box[2] - title_box[0])) / 2, 48), title,
              font=title_font, fill='#18324A')
    subtitle = 'Shared prefixes appear once • Green double-border nodes mark complete words'
    subtitle_box = draw.textbbox((0, 0), subtitle, font=note_font)
    draw.text(((width - (subtitle_box[2] - subtitle_box[0])) / 2, 122), subtitle,
              font=note_font, fill='#586574')

    def draw_edges(node):
        for child in node.children.values():
            draw.line((node.x, node.y, child.x, child.y), fill='#9AA8B6', width=5)
            draw_edges(child)

    draw_edges(root)

    radius = 39

    def draw_nodes(node):
        left, top = node.x - radius, node.y - radius
        right, bottom = node.x + radius, node.y + radius
        fill = GREEN if node.terminal else '#EAF1F7'
        edge = GREEN_EDGE if node.terminal else '#355E7C'
        draw.ellipse((left, top, right, bottom), fill=fill, outline=edge, width=6)
        if node.terminal:
            draw.ellipse((left + 8, top + 8, right - 8, bottom - 8), outline=edge, width=3)
        label = node.char if node is not root else 'R'
        label_font = bold if node is root else regular
        box = draw.textbbox((0, 0), label, font=label_font)
        tx = node.x - (box[2] - box[0]) / 2
        ty = node.y - (box[3] - box[1]) / 2 - 3
        draw.text((tx, ty), label, font=label_font, fill=INK)
        for child in node.children.values():
            draw_nodes(child)

    draw_nodes(root)

    legend_y = 1590
    draw.rounded_rectangle((300, legend_y, 2900, 1695), radius=24,
                           fill='#F4F7FA', outline='#CCD6DF', width=3)
    legend = ('Start at R and follow one letter per node. A green double-border node means '
              'the letters from R form a complete word; that node may still have children.')
    box = draw.textbbox((0, 0), legend, font=note_font)
    draw.text(((width - (box[2] - box[0])) / 2, legend_y + 34), legend,
              font=note_font, fill='#334455')

    image.save(PNG_PATH, dpi=(300, 300))


def set_cell_or_run_font(run, size=11, bold=False, italic=False, color='000000'):
    run.font.name = 'Calibri'
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def configure_page(section, landscape=False):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_cell_or_run_font(run, size=10, color=MUTED)
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, separate, end])


def configure_header_footer(section):
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ''
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run('CS 508  |  Module Seven Journal')
    set_cell_or_run_font(r, size=9, bold=True, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ''
    add_page_number(fp)


def set_spacing(paragraph, before=0, after=6, line=1.10):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=4, after=4, line=1.0)
    r = p.add_run(text)
    set_cell_or_run_font(r, size=23, bold=True, color='000000')
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p2, before=0, after=12, line=1.0)
        r2 = p2.add_run(subtitle)
        set_cell_or_run_font(r2, size=11, color=MUTED)


def add_heading(doc, text, level=1, page_break_before=False):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = page_break_before
    if level == 1:
        set_spacing(p, before=16, after=8, line=1.0)
        r = p.add_run(text)
        set_cell_or_run_font(r, size=16, bold=True, color=BLUE)
    else:
        set_spacing(p, before=12, after=6, line=1.0)
        r = p.add_run(text)
        set_cell_or_run_font(r, size=13, bold=True, color=BLUE)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=6, line=1.10)
    r = p.add_run(text)
    set_cell_or_run_font(r, size=11)
    return p


def add_bullet(doc, label, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_spacing(p, before=0, after=8, line=1.167)
    r1 = p.add_run(label + ': ')
    set_cell_or_run_font(r1, size=11, bold=True, color=DARK_BLUE)
    r2 = p.add_run(text)
    set_cell_or_run_font(r2, size=11)


def add_reference(doc, parts):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    set_spacing(p, before=0, after=6, line=1.10)
    for text, italic in parts:
        r = p.add_run(text)
        set_cell_or_run_font(r, size=11, italic=italic)


def build_document(node_count):
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    first = doc.sections[0]
    configure_page(first, landscape=True)
    configure_header_footer(first)

    add_title(doc, 'Prefix Tree (Trie) Diagram', 'CS 508  |  August 19, 2026')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=4, line=1.0)
    p.add_run().add_picture(str(PNG_PATH), width=Inches(8.55))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(cap, before=0, after=4, line=1.0)
    rc = cap.add_run(f'Figure 1. Minimal standard trie: {node_count} nodes including the root; no extra end-of-word nodes.')
    set_cell_or_run_font(rc, size=9, italic=True, color=MUTED)
    words_p = doc.add_paragraph()
    words_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(words_p, before=0, after=0, line=1.0)
    rw1 = words_p.add_run('Words represented: ')
    set_cell_or_run_font(rw1, size=8.5, bold=True, color=DARK_BLUE)
    rw2 = words_p.add_run(', '.join(sorted(WORDS)))
    set_cell_or_run_font(rw2, size=8.5, color=MUTED)

    second = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_page(second, landscape=False)
    second.header.is_linked_to_previous = False
    second.footer.is_linked_to_previous = False
    configure_header_footer(second)

    add_title(doc, 'Trie Structure and Strengths')
    add_heading(doc, 'How the Diagram Uses the Fewest Nodes')
    add_body(doc, 'The diagram starts at the root, marked R, and stores one letter in each node. A complete word is found by following a path from the root. For example, c → a → r forms “car.” The r node is marked as a complete word, but the same node continues to b, d, e, p, r, and t for words such as “carb,” “card,” “care,” “carp,” “carry,” and “cart.” This is what allows one trie to store related words without rebuilding the shared c-a-r prefix each time.')
    add_body(doc, f'The 25 words require {node_count} total nodes when the root is counted. The diagram does not create separate end-of-word nodes. Instead, a green double border marks the existing letter node where a word ends. This keeps the node count low while still showing words that end inside another word’s path. For example, “car” ends at r even though longer words continue below it. Every shared prefix appears only once.')

    add_heading(doc, 'Strengths Compared With Other Data Structures')
    add_body(doc, 'The biggest strength of a trie is that lookup depends mainly on the length of the word or prefix instead of the total number of stored words. If a search contains m letters, the trie follows at most m character links. This makes a trie especially useful when the question is not only “Does this exact word exist?” but also “What words begin with these letters?” (Goodrich et al., 2014).')
    add_bullet(doc, 'Compared with an unsorted array or linked list', 'An array or linked list may require checking many complete words one at a time. A trie follows the letters of the search directly. A linked list also has to move node by node and does not provide direct access to a prefix branch.')
    add_bullet(doc, 'Compared with a sorted array', 'Binary search can locate an exact word efficiently, but autocomplete still requires finding the prefix boundary and scanning nearby entries. A trie reaches the prefix node directly, and every descendant below that node is a possible completion.')
    add_bullet(doc, 'Compared with a binary search tree', 'A balanced BST can search in O(log n) comparisons, but each comparison may examine several characters. A trie makes one decision for each character and naturally groups every word with the same prefix.')
    add_bullet(doc, 'Compared with a hash table', 'A hash table is usually very fast for exact-word lookup, but it does not naturally preserve alphabetical order or shared prefixes. Finding every word beginning with “car” may require scanning many keys. In a trie, the c-a-r path immediately identifies the correct branch.')

    add_heading(doc, 'Where a Trie Is Most Useful', page_break_before=True)
    add_body(doc, 'A trie is a strong choice for autocomplete, spell-checking, dictionaries, contact searches, command suggestions, and routing systems that match text by prefix. It can return all completions below a prefix node, and alphabetical output is straightforward when child links are visited in letter order. Shared prefixes can also reduce repeated storage when many words begin with the same letters.')
    add_body(doc, 'The tradeoff is that each trie node needs child references and an end-of-word marker. If the words share very few prefixes, that overhead may use more memory than a compact array or hash table. For this assignment, however, many words share “car,” “ca,” “co,” “la,” and other prefixes, so the trie gives a clear and efficient representation.')

    add_heading(doc, 'References')
    add_reference(doc, [
        ('Goodrich, M. T., Tamassia, R., & Goldwasser, M. H. (2014). ', False),
        ('Data structures and algorithms in Java', True),
        (' (6th ed.). Wiley.', False),
    ])
    add_reference(doc, [
        ('OpenAI. (2026). ', False),
        ('ChatGPT', True),
        (' [Large language model]. https://chatgpt.com/', False),
    ])

    add_heading(doc, 'AI Use Disclosure')
    add_body(doc, 'I used ChatGPT to help organize the given words into a trie, verify that every word was represented, create the diagram, and improve the explanation and document formatting. I reviewed the final trie paths, node count, and written explanation before submission.')

    doc.save(DOCX_PATH)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    root = build_trie(WORDS)
    node_count = count_nodes(root)
    assert len(WORDS) == 25
    assert node_count == 45
    create_diagram(root)
    build_document(node_count)
    print(f'Created {PNG_PATH}')
    print(f'Created {DOCX_PATH}')
    print(f'Words: {len(WORDS)}; nodes including root: {node_count}')


if __name__ == '__main__':
    main()
