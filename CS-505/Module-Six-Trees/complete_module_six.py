from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK

BASE_DIR = Path(__file__).resolve().parent
SOURCE = BASE_DIR / 'input' / 'CS 505 Module Six Activity Template.docx'
OUTPUT = BASE_DIR / 'output' / 'CS 505 Module Six Activity Completed.docx'
OUTPUT.parent.mkdir(parents=True, exist_ok=True)


def set_font(run, name='Calibri', size=11, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        if child.tag != qn('w:pPr'):
            p.remove(child)


def clear_text_preserve_drawings(paragraph):
    for run in paragraph._element.findall('.//' + qn('w:r')):
        has_drawing = run.find('.//' + qn('w:drawing')) is not None
        has_pict = run.find('.//' + qn('w:pict')) is not None
        for text_node in run.findall('.//' + qn('w:t')):
            text_node.text = ''
        if not has_drawing and not has_pict:
            parent = run.getparent()
            if parent is not None:
                parent.remove(run)


def write_answer(paragraph, text, size=11, preserve_drawings=False):
    if preserve_drawings:
        clear_text_preserve_drawings(paragraph)
    else:
        clear_paragraph(paragraph)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    run = paragraph.add_run(text)
    set_font(run, size=size)


def replace_placeholder(paragraph, answer):
    full = paragraph.text
    if '[Insert text.]' not in full:
        raise ValueError(f'No placeholder in: {full!r}')
    question = full.replace('\n[Insert text.]', '').replace('[Insert text.]', '').rstrip()
    clear_paragraph(paragraph)
    q = paragraph.add_run(question)
    set_font(q, size=11)
    q.add_break()
    a = paragraph.add_run(answer)
    set_font(a, size=11)
    paragraph.paragraph_format.space_after = Pt(8)


def insert_after(paragraph, text, style='Normal', font='Calibri', size=11,
                 left_indent=0.5, space_after=8):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_p.addnext(new_para._p)
    new_para.style = style
    new_para.paragraph_format.left_indent = Inches(left_indent)
    new_para.paragraph_format.space_before = Pt(2)
    new_para.paragraph_format.space_after = Pt(space_after)
    r = new_para.add_run(text)
    set_font(r, name=font, size=size)
    return new_para


doc = Document(SOURCE)

# Part One placeholders are separate paragraphs after each graph.
part_one_answers = [
    ('Graph 1 is a tree because all five nodes are connected and there is no cycle. '
     'Using node 0 as the root, the leaves are nodes 2 and 4. The internal nodes are '
     '1 and 3. Node 0 is the root. A different node could technically be chosen as '
     'the root because the graph itself is undirected, but node 0 gives a clear rooted view.'),
    ('Graph 2 is not a tree because it contains a cycle: 0 -> 1 -> 3 -> 4 -> 0. '
     'A tree cannot contain any cycle.'),
    ('Graph 3 is not a tree because it is disconnected. Node 4 is isolated from nodes '
     '0, 1, 2, and 3. A tree must connect every node.')
]
placeholder_indices = [i for i, p in enumerate(doc.paragraphs)
                       if p.text.strip() == '[Insert text.]']
for index, answer in zip(placeholder_indices[:3], part_one_answers):
    write_answer(doc.paragraphs[index], answer, preserve_drawings=True)

# Parts Two and Three: question and placeholder share one paragraph.
answers = {
    'Prove that a tree with n vertices has n − 1 edges.':
        'Start with one vertex. It has 0 edges, which equals 1 - 1. Each time a new '
        'vertex is added to a tree, exactly one edge must connect it to the existing '
        'tree. Adding more than one would create a cycle, and adding none would leave '
        'it disconnected. Therefore, after adding n - 1 more vertices, the tree has '
        'n - 1 edges.',
    'Prove that a tree with n vertices has at least one leaf.':
        'Take the longest path in the tree. An endpoint of that path cannot have another '
        'unused neighbor, because then the path could be made longer. Therefore, each '
        'endpoint has degree 1 and is a leaf. For a one-node tree, the root is also a '
        'leaf. Thus, every finite tree has at least one leaf.',
    'Explain why there is exactly one path between any two nodes in a tree.':
        'A tree is connected, so at least one path exists between any two nodes. If two '
        'different paths existed, following one path out and the other path back would '
        'form a cycle. Trees cannot contain cycles, so only one path can exist.',
    'Describe an algorithm to determine if a graph contains a cycle.':
        'Use depth-first search (DFS). Mark each node when it is visited. For an '
        'undirected graph, remember the node that led to the current node, called the '
        'parent. If DFS reaches a visited neighbor that is not the parent, a cycle '
        'exists. Run DFS from every unvisited node so disconnected parts are also checked.',
    'Explain pre-order traversal, including an example of a pre-order traversal of a binary tree.':
        'Pre-order visits Root, Left, Right. For the provided tree, visit 1 first, then '
        'the left subtree 2, 4, 5, and then the right subtree 3, 6. The result is '
        '1, 2, 4, 5, 3, 6.',
    'Explain in-order traversal, including an example of an in-order traversal of a binary tree.':
        'In-order visits Left, Root, Right. For the provided tree, the result is '
        '4, 2, 5, 1, 3, 6.',
    'Explain post-order traversal, including an example of a post-order traversal of a binary tree.':
        'Post-order visits Left, Right, Root. For the provided tree, the result is '
        '4, 5, 2, 6, 3, 1.',
}

for p in doc.paragraphs:
    for question, answer in answers.items():
        if p.text.startswith(question) and '[Insert text.]' in p.text:
            replace_placeholder(p, answer)
            break

# Remaining Part Three answer-only placeholders in question order.
remaining = [p for p in doc.paragraphs if p.text.strip() == '[Insert text.]']
remaining_answers = [
    'Pre-order: 1, 2, 4, 5, 3, 6\n'
    'In-order: 4, 2, 5, 1, 3, 6\n'
    'Post-order: 4, 5, 2, 6, 3, 1',
    'The height is 2 when height is counted by edges from the root to the deepest '
    'leaf. The tree has 3 levels if levels are counted instead.',
    'There are 3 leaf nodes: 4, 5, and 6. A leaf has no children.',
    'Node 5 is a leaf, so it can simply be removed. Node 2 then has only its left '
    'child, node 4. The remaining tree is:\n'
    '        1\n'
    '      /   \\\n'
    '     2     3\n'
    '    /       \\\n'
    '   4         6',
    'Node 3 already has node 6 as its right child. A binary-tree node cannot have two '
    'right children, so this insertion is not valid unless node 6 is moved or removed. '
    'If the instruction means to replace node 6 with node 7, the result is:\n'
    '        1\n'
    '      /   \\\n'
    '     2     3\n'
    '    / \\     \\\n'
    '   4   5     7'
]
for p, answer in zip(remaining, remaining_answers):
    write_answer(p, answer, size=10.5)
    if '\n' in answer:
        for run in p.runs:
            set_font(run, name='Consolas', size=9.5)

# Part Four had no response placeholders. Insert answers directly after prompts.
part_four = {
    'Explain how a file system can be represented as a tree.':
        'A file system is hierarchical. The main drive or top folder is at the top. '
        'Folders branch into subfolders and files. Each file or folder has one parent '
        'location, except the root, and a path shows how to travel from the root to an item.',
    'Identify the root, internal nodes, and leaf nodes in a file system tree.':
        'The root is the drive or top folder, such as C:\\ or /. Internal nodes are '
        'folders that contain other folders or files. Leaf nodes are usually files or '
        'empty folders because they have no children.',
    'Describe how decision trees can be used to make decisions, including an example of a decision tree for a simple classification problem.':
        'A decision tree asks one question at each internal node and follows a branch '
        'based on the answer. A leaf gives the final class. Example: to classify an '
        'email, ask "Does it contain suspicious links?" If no, classify it as Normal. '
        'If yes, ask "Is the sender known?" A known sender may be Review, while an '
        'unknown sender may be Spam.',
    'Explain the properties of a binary search tree.':
        'A binary search tree (BST) has at most two children per node. Every value in a '
        'node\'s left subtree is smaller than the node, and every value in its right '
        'subtree is larger. The left and right subtrees must also follow the same rule. '
        'An in-order traversal returns the values in sorted order. Search, insertion, '
        'and deletion take O(h) time, where h is the tree height.',
    'Implement a binary search tree using pseudocode.':
        'SET root to NULL\n\n'
        'DEFINE NODE(value)\n'
        '    SET node.value to value\n'
        '    SET node.left to NULL\n'
        '    SET node.right to NULL\n'
        '    RETURN node\n'
        'END DEFINE\n\n'
        'FUNCTION INSERT(node, value)\n'
        '    IF node is NULL\n'
        '        RETURN NODE(value)\n'
        '    END IF\n'
        '    IF value is less than node.value\n'
        '        SET node.left to INSERT(node.left, value)\n'
        '    ELSE IF value is greater than node.value\n'
        '        SET node.right to INSERT(node.right, value)\n'
        '    END IF\n'
        '    RETURN node\n'
        'END FUNCTION\n\n'
        'FUNCTION SEARCH(node, target)\n'
        '    WHILE node is not NULL\n'
        '        IF target equals node.value\n'
        '            RETURN node\n'
        '        ELSE IF target is less than node.value\n'
        '            SET node to node.left\n'
        '        ELSE\n'
        '            SET node to node.right\n'
        '        END IF\n'
        '    END WHILE\n'
        '    RETURN NOT FOUND\n'
        'END FUNCTION'
}

# Insert in reverse document order so positions stay stable.
targets = []
for p in doc.paragraphs:
    if p.text in part_four:
        targets.append(p)
for p in reversed(targets):
    answer = part_four[p.text]
    new_p = insert_after(p, answer, size=10.5 if '\n' not in answer else 9.5)
    if '\n' in answer:
        for run in new_p.runs:
            set_font(run, name='Consolas', size=9.5)

doc.core_properties.title = 'CS 505 Module Six Activity Completed'
doc.core_properties.author = 'Student'
doc.save(OUTPUT)
print(OUTPUT)
