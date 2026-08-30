from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT = "outputs/Combinatorics_Questions_1_to_12.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GRAY = "555555"


def font(run, size=11, bold=False, italic=False, color="000000"):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def paragraph(doc, text="", bold=False, italic=False, color="000000", after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    font(r, bold=bold, italic=italic, color=color)
    return p


def question(doc, number, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{number}. {text}")
    font(r, size=16, bold=True, color=BLUE)


def label(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    font(r, size=12, bold=True, color=DARK_BLUE)


def formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    font(r, size=12, bold=True, color=DARK_BLUE)


def answer(doc, text):
    label(doc, "Final Answer")
    paragraph(doc, text, bold=True, after=8)


def page_number(paragraph_obj):
    paragraph_obj.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph_obj.add_run("Page ")
    font(r, size=9, color="777777")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph_obj._p.append(fld)


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

header = sec.header.paragraphs[0]
header.text = "CS 505 Module Five Activity | Combinatorics"
font(header.runs[0], size=9, color="777777")
page_number(sec.footer.paragraphs[0])

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(6)
r = title.add_run("Combinatorics Problems")
font(r, size=24, bold=True, color=DARK_BLUE)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(18)
r = sub.add_run("Questions 1-12: Methods, Proofs, and Simple Explanations")
font(r, size=12, color=GRAY)

# 1
question(doc, 1, "Explain the number of different ways 10 people can be arranged in a line.")
label(doc, "Method")
paragraph(doc, "I used a permutation and factorial because the order of the people matters.")
label(doc, "Proof")
paragraph(doc, "There are 10 choices for the first spot, 9 for the second spot, and so on until one person remains.")
formula(doc, "10! = 10 x 9 x 8 x 7 x 6 x 5 x 4 x 3 x 2 x 1 = 3,628,800")
answer(doc, "The 10 people can be arranged in 3,628,800 different ways.")

# 2
question(doc, 2, "Explain the number of different ways a committee of five people can be selected from a group of 20 people.")
label(doc, "Method")
paragraph(doc, "I used a combination because the order of the five committee members does not matter.")
label(doc, "Proof")
formula(doc, "C(20, 5) = 20! / (5!15!)")
formula(doc, "= (20 x 19 x 18 x 17 x 16) / (5 x 4 x 3 x 2 x 1) = 15,504")
answer(doc, "There are 15,504 different committees.")

doc.add_page_break()

# 3
question(doc, 3, "Explain the number of different ways two cards can be selected from a standard deck of 52 cards.")
label(doc, "Method")
paragraph(doc, "I used a combination because selecting Card A and then Card B is the same pair as selecting Card B and then Card A.")
label(doc, "Proof")
formula(doc, "C(52, 2) = 52! / (2!50!) = (52 x 51) / 2 = 1,326")
answer(doc, "There are 1,326 different two-card selections.")

# 4
question(doc, 4, "Explain how many different 6-letter words can be formed using the letters of the word MISSISSIPPI.")
label(doc, "Method")
paragraph(doc, "I used multiset permutations because MISSISSIPPI contains repeated letters: M appears 1 time, I appears 4 times, S appears 4 times, and P appears 2 times. A 6-letter word cannot use a letter more times than it appears.")
label(doc, "Proof")
paragraph(doc, "For every allowed choice of 6 letters, I counted its arrangements with 6!/(m!i!s!p!), then added the results. The allowed cases with no M total 470 arrangements. The allowed cases with one M total 1,140 arrangements.")
formula(doc, "470 + 1,140 = 1,610")
answer(doc, "There are 1,610 different 6-letter arrangements. These are letter arrangements and do not have to be dictionary words.")

doc.add_page_break()

# 5
question(doc, 5, "Describe real-world applications of combinations and permutations.")
label(doc, "Explanation")
paragraph(doc, "Permutations are used when order matters. Examples include planning seating arrangements, creating schedules, ranking race results, and testing possible passwords. Combinations are used when order does not matter. Examples include choosing a committee, selecting a sports team, picking lottery numbers, and choosing products for a sample.")
answer(doc, "Use permutations for ordered arrangements and combinations for groups or selections.")

# 6
question(doc, 6, "Explain the difference between permutations and combinations.")
label(doc, "Explanation")
paragraph(doc, "The main question is: Does changing the order create a different result?")
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
widths = [Inches(1.45), Inches(2.45), Inches(2.60)]
for idx, cell in enumerate(table.rows[0].cells):
    cell.width = widths[idx]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ["Method", "When to Use It", "Example"][idx]
    for run in cell.paragraphs[0].runs:
        font(run, bold=True, color=DARK_BLUE)
for values in [
    ("Permutation", "Order matters", "ABC and BAC are different"),
    ("Combination", "Order does not matter", "Choosing A and B is the same group as B and A"),
]:
    cells = table.add_row().cells
    for idx, value in enumerate(values):
        cells[idx].width = widths[idx]
        cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[idx].text = value
        for run in cells[idx].paragraphs[0].runs:
            font(run)
paragraph(doc, "", after=2)
answer(doc, "A permutation counts arrangements, while a combination counts selections.")

# 7
doc.add_page_break()

# 7
question(doc, 7, "Explain the number of different ways six identical balls can be distributed into six distinct boxes.")
label(doc, "Method")
paragraph(doc, "I used stars and bars. The six identical balls are the stars, and five bars separate the balls into six different boxes. Empty boxes are allowed.")
label(doc, "Proof")
paragraph(doc, "There are 11 total symbols: 6 stars and 5 bars. Choose where the 5 bars go.")
formula(doc, "C(6 + 6 - 1, 6 - 1) = C(11, 5) = 462")
answer(doc, "The balls can be distributed in 462 different ways.")

# 8
question(doc, 8, "Explain the number of different ways 15 identical candies can be distributed to 3 children, if each child must receive at least three candies.")
label(doc, "Method")
paragraph(doc, "First, give each child 3 candies. This uses 9 candies and guarantees the minimum. Then use stars and bars to distribute the 6 candies that remain.")
label(doc, "Proof")
formula(doc, "15 - (3 x 3) = 6 candies remaining")
formula(doc, "C(6 + 3 - 1, 3 - 1) = C(8, 2) = 28")
answer(doc, "The candies can be distributed in 28 different ways.")

doc.add_page_break()

# 9
question(doc, 9, "Explain how many non-negative integer solutions there are to x1 + x2 + x3 = 20.")
label(doc, "Method")
paragraph(doc, "I used stars and bars because the variables may be zero or greater. Think of 20 stars separated into 3 groups by 2 bars.")
label(doc, "Proof")
formula(doc, "C(20 + 3 - 1, 3 - 1) = C(22, 2) = 231")
answer(doc, "There are 231 non-negative integer solutions.")

# 10
question(doc, 10, "Explain how many positive integer solutions there are to x1 + x2 + x3 = 10.")
label(doc, "Method")
paragraph(doc, "Positive means each variable must be at least 1. Give 1 to each variable first. This leaves 7 to distribute without restrictions.")
label(doc, "Proof")
formula(doc, "10 - 3 = 7")
formula(doc, "C(7 + 3 - 1, 3 - 1) = C(9, 2) = 36")
answer(doc, "There are 36 positive integer solutions.")

doc.add_page_break()

# 11
question(doc, 11, "Describe how the stars and bars method can be adjusted to handle problems with constraints.")
label(doc, "Explanation")
paragraph(doc, "For a minimum requirement, give each person or box the required minimum first. Then use stars and bars on what remains. For example, if each child needs at least 3 candies, give each child 3 before distributing the rest.")
paragraph(doc, "For a positive-integer problem, give each variable 1 first. For a maximum limit, count all distributions and subtract the ones that break the limit, often using inclusion-exclusion.")
answer(doc, "Handle minimums by assigning them first; handle maximums by removing invalid cases.")

# 12
question(doc, 12, "Describe real-world applications of stars and bars.")
label(doc, "Explanation")
paragraph(doc, "Stars and bars can be used whenever identical items are divided among different people, locations, or categories. A business can divide identical products among stores, a teacher can divide identical supplies among classrooms, and a company can assign whole budget units among departments.")
paragraph(doc, "It can also count possible workloads, such as distributing identical service requests among teams, as long as only the number assigned to each team matters.")
answer(doc, "Stars and bars helps count ways to distribute identical items among distinct groups.")

doc.core_properties.title = "Combinatorics Problems: Questions 1-12"
doc.core_properties.subject = "Methods, proofs, and explanations"
doc.core_properties.author = "Student"
doc.save(OUTPUT)
print(OUTPUT)
