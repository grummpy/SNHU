from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


OUTPUT = "outputs/Committee_Selection_Problem.docx"


def set_font(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, size=16, bold=True, color="2E74B5")
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(6)
r = title.add_run("Committee Selection Problem")
set_font(r, size=22, bold=True, color="1F4D78")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(18)
r = subtitle.add_run("Combinatorics: Using a Combination")
set_font(r, size=12, color="555555")

add_heading(doc, "Problem")
add_body(doc, "Explain the number of different ways a committee of five people can be selected from a group of 20 people.")

add_heading(doc, "Method")
add_body(doc, "I used a combination because the order of the five people does not matter. We are only choosing which people will be on the committee.")

formula = doc.add_paragraph()
formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
formula.paragraph_format.space_before = Pt(8)
formula.paragraph_format.space_after = Pt(8)
r = formula.add_run("C(20, 5) = 20! / [5!(20 - 5)!]")
set_font(r, size=13, bold=True, color="1F4D78")

add_heading(doc, "Proof")
add_body(doc, "First, substitute the numbers into the combination formula:")

for equation in [
    "C(20, 5) = 20! / (5!15!)",
    "C(20, 5) = (20 x 19 x 18 x 17 x 16) / (5 x 4 x 3 x 2 x 1)",
    "C(20, 5) = 1,860,480 / 120",
    "C(20, 5) = 15,504",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(equation)
    set_font(r, size=12)

add_heading(doc, "Final Answer")
p = add_body(doc, "There are 15,504 different ways to select a committee of five people from a group of 20 people.")
for run in p.runs:
    run.bold = True

doc.core_properties.title = "Committee Selection Problem"
doc.core_properties.subject = "Combination method and proof"
doc.core_properties.author = "Student"
doc.save(OUTPUT)
print(OUTPUT)
