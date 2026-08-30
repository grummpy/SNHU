from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT = Path('discrete_modeling_paper/Discrete Modeling Real-World Problems.docx')


def font(run, size=12, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    rpr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    r = p.add_run(text)
    font(r)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    font(r, bold=True)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    field = OxmlElement('w:fldSimple')
    field.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(field)


def reference(doc, text):
    p = body(doc, text, indent=False)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.bottom_margin = Inches(1)
section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.5)

normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 2
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)
add_page_number(section.header.paragraphs[0])

# APA-style student title page.
p = doc.add_paragraph()
p.paragraph_format.space_before = Inches(2)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 2
r = p.add_run('Using Discrete Models to Solve Real-World Computing Problems')
font(r, bold=True)

for line in [
    'Student Name',
    'Southern New Hampshire University',
    'CS 505',
    'Instructor Name',
    'August 2026',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2
    r = p.add_run(line)
    font(r)

doc.add_page_break()

heading(doc, 'Using Discrete Models to Solve Real-World Computing Problems')
body(doc,
    'Discrete models describe systems as separate objects, states, or events instead '
    'of treating everything as continuously changing. They are useful in computing '
    'because many real problems involve individual roads, school terms, customer '
    'requests, or equipment failures. However, one model is not best for every '
    'problem. This paper matches graph theory to a daily travel problem, a Markov '
    'chain to an education problem, and discrete-event simulation to a career '
    'problem involving utility-company operations.')

heading(doc, 'Daily Life: Finding an Efficient Route', 2)
body(doc,
    'A common daily computing problem is choosing the best route between home, work, '
    'school, stores, and other stops. A navigation system must consider many '
    'intersections and roads. Each road may have a different distance, estimated '
    'travel time, toll, or traffic level. The challenge is to choose a useful route '
    'without testing every possible sequence of roads. The route may also need to '
    'change when traffic or a road closure changes the available choices.')
body(doc,
    'Graph theory is an appropriate model for this problem. Intersections or '
    'destinations become vertices, and roads become edges. Each edge can receive a '
    'weight such as minutes or miles. A shortest-path algorithm then finds a path '
    'with the smallest total weight. If traffic conditions change, the program can '
    'update edge weights and calculate the path again. Gajjar et al. (2022) describe '
    'shortest-path reconfiguration and identify road-network changes and data '
    'rerouting as natural graph applications.')
body(doc,
    'The main advantage of graph theory is that it matches the physical road system '
    'directly. It supports efficient algorithms, allows different road costs, and can '
    'show alternative routes. Its disadvantage is that the answer depends on the '
    'quality of the graph and its weights. A supposedly shortest route may be poor if '
    'traffic data are old, a road is missing, or the model ignores parking and '
    'personal preferences. Very large changing networks also require frequent data '
    'updates and repeated calculations.')
body(doc,
    'Graph theory was selected because the problem is mainly about connections and '
    'path cost. Roads and intersections are already separate, countable objects, so '
    'they fit a discrete graph. The model also respects the problem constraints: a '
    'route may use only existing roads, road closures can be removed as edges, and '
    'time or distance limits can be stored as weights. A Markov chain would focus on '
    'probabilities rather than finding a path, while a simulation would be more '
    'complicated than necessary for one person seeking a route.')

heading(doc, 'Education: Predicting Student Progress', 2)
body(doc,
    'An education-related computing problem is identifying whether students are '
    'likely to continue, repeat a level, graduate, or leave a program. Schools want '
    'to provide help early, but student progress occurs over several terms and does '
    'not follow one guaranteed path. The available data may also be limited for small '
    'student groups, creating uncertainty in ordinary graduation estimates.')
body(doc,
    'A Markov chain can represent the process with states such as first year, second '
    'year, third year, graduated, and withdrawn. Historical records provide '
    'transition probabilities, such as the chance that a first-year student advances '
    'or withdraws during the next term. Multiplying the current distribution of '
    'students by the transition matrix estimates future distributions. Graduation '
    'and withdrawal can be absorbing states because students do not return to an '
    'active academic level in the basic model. Tedeschi et al. (2023) found that a '
    'Markov model increased confidence and reduced bias in graduation-rate estimates '
    'for small underrepresented and first-generation student groups.')
body(doc,
    'The advantages are that a Markov chain is compact, uses understandable states, '
    'and supports probability-based forecasts across several semesters. It can also '
    'help a school compare outcomes before and after an academic support program. The '
    'main disadvantage is the Markov assumption: the next state is predicted mainly '
    'from the current state. Real performance can also depend on finances, work, '
    'health, course difficulty, and earlier experiences. Transition probabilities '
    'based on old data may become inaccurate when policies or student populations '
    'change. The results should guide support, not make final decisions about one '
    'student.')
body(doc,
    'The Markov chain was chosen because the problem is about movement among a small '
    'set of academic states over equal time steps. It handles uncertainty better than '
    'a basic graph that only shows which transitions are possible. It also fits the '
    'constraints of term-by-term school records and incomplete certainty about '
    'individual outcomes. The states and probabilities can be updated each year, '
    'making the model useful while acknowledging that it is an estimate rather than '
    'a promise.')

heading(doc, 'Career: Scheduling Utility-Company Work', 2)
body(doc,
    'A career-related computing problem is planning crews for a utility company. '
    'Normal service orders, emergency outages, travel, repairs, shift changes, and '
    'equipment availability all affect the schedule. Requests arrive at different '
    'times, repair durations vary, and a major storm can create a sudden queue. The '
    'company needs enough workers for acceptable response times, but keeping too many '
    'workers idle is expensive.')
body(doc,
    'Discrete-event simulation is a strong model for this problem. The simulation '
    'clock moves from one important event to the next, such as a service request '
    'arriving, a crew becoming available, a repair beginning, or a repair ending. '
    'Inputs can include arrival rates, travel times, repair-time distributions, crew '
    'skills, priorities, and equipment failures. The utility can run many scenarios '
    'with different crew levels and dispatch rules, then compare average wait time, '
    'queue length, crew use, cost, and the number of late jobs. Tiacci and Rossi '
    '(2024) explain that discrete-event simulation is important for testing dynamic '
    'schedules affected by events such as new jobs and machine failures.')
body(doc,
    'The advantage of discrete-event simulation is safe experimentation. The company '
    'can test a storm plan or staffing change without disrupting real customers. It '
    'can represent queues, limited workers, priorities, randomness, and interactions '
    'that a simple average may hide. The disadvantages are that a detailed model '
    'takes time to build and validate, and it may require a large amount of accurate '
    'data. A simulation does not automatically find the perfect plan; it compares the '
    'scenarios that the analyst chooses to run. Poor assumptions can produce '
    'realistic-looking but misleading results.')
body(doc,
    'Discrete-event simulation was selected because timing, queues, limited resources, '
    'and random events are the main requirements of utility scheduling. A graph can '
    'model the road or power network, but it cannot by itself show how a queue changes '
    'during a shift. A Markov chain can model broad system states, but it normally '
    'does not preserve the detailed timing and resource competition needed here. '
    'Simulation directly addresses constraints such as crew limits, required skills, '
    'job priority, travel time, and uncertain repair duration.')

heading(doc, 'Conclusion', 2)
body(doc,
    'The best discrete model depends on the structure of the problem. Graph theory '
    'fits daily route planning because it represents connections and weighted paths. '
    'A Markov chain fits student progress because it represents uncertain movement '
    'between academic states over time. Discrete-event simulation fits utility '
    'scheduling because separate arrivals, completions, queues, and resources change '
    'the system. Each model simplifies reality and has limitations, but each provides '
    'a practical way to organize information, test choices, and support better '
    'computing decisions.')

heading(doc, 'References')
reference(doc,
    'Gajjar, K., Jha, A. V., Kumar, M., & Lahiri, A. (2022). Reconfiguring shortest '
    'paths in graphs. Proceedings of the AAAI Conference on Artificial Intelligence, '
    '36(9), 9758-9766. https://doi.org/10.1609/aaai.v36i9.21211')
reference(doc,
    'Tedeschi, M. N., Hose, T. M., Mehlman, E. K., Franklin, S., & Wong, T. E. '
    '(2023). Improving models for student retention and graduation using Markov '
    'chains. PLOS ONE, 18(6), e0287775. '
    'https://doi.org/10.1371/journal.pone.0287775')
reference(doc,
    'Tiacci, L., & Rossi, A. (2024). A discrete event simulator to implement deep '
    'reinforcement learning for the dynamic flexible job shop scheduling problem. '
    'Simulation Modelling Practice and Theory, 134, 102948. '
    'https://doi.org/10.1016/j.simpat.2024.102948')

heading(doc, 'AI Use Disclosure')
body(doc,
    'I used ChatGPT to help identify suitable discrete modeling techniques, organize '
    'the paper, simplify explanations, and format the references. I reviewed the '
    'content and am responsible for the final submission.', indent=False)
reference(doc,
    'OpenAI. (2026). ChatGPT [Large language model]. https://chatgpt.com/')

doc.core_properties.title = 'Using Discrete Models to Solve Real-World Computing Problems'
doc.core_properties.author = 'Student'
doc.save(OUTPUT)
print(OUTPUT)
