from pathlib import Path
from itertools import product
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
INPUT_DIR = ROOT / 'input'
OUTPUT_DIR = ROOT / 'output'
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
BOOL_SOURCE = INPUT_DIR / 'CS 505 Module Seven Activity Boolean Algebra Template.docx'
ADJ_SOURCE = INPUT_DIR / 'CS 505 Module Seven Activity Adjacency Matrix Template.docx'
BOOL_OUT = OUTPUT_DIR / 'CS 505 Module Seven Boolean Algebra Completed.docx'
ADJ_OUT = OUTPUT_DIR / 'CS 505 Module Seven Adjacency Matrix Completed.docx'
GUIDE_OUT = OUTPUT_DIR / 'CS 505 Module Seven Simple Explainer Tables.docx'
CIRCUIT_PATH = OUTPUT_DIR / 'boolean_logic_circuit.png'

BLUE = '1F4E79'
LIGHT_BLUE = 'D9EAF7'
LIGHT_GRAY = 'F2F4F7'
GREEN = 'E2F0D9'


def set_run_font(run, name='Times New Roman', size=12, bold=False,
                 italic=False, color='000000'):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn('w:ascii'), name)
    rpr.rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def replace_placeholder(paragraph, answer):
    full = paragraph.text.replace('[Insert text.]', answer)
    paragraph.clear()
    run = paragraph.add_run(full)
    set_run_font(run, size=12)
    paragraph.paragraph_format.space_after = Pt(6)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width_dxa))
    tc_w.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), str(indent))
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])


def format_table(table, font_size=10.5, header_fill=LIGHT_BLUE,
                 widths=None, indent=120):
    try:
        table.style = 'Table Grid'
    except KeyError:
        # Some SNHU templates do not include Word's built-in Table Grid style.
        # Add the same simple grid directly so the source style set is preserved.
        tbl_pr = table._tbl.tblPr
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '6')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'B7C3CC')
            borders.append(border)
        tbl_pr.append(borders)
    if widths:
        set_table_geometry(table, widths, indent)
    for r_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_index == 0:
                shade_cell(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(run, size=font_size,
                                 bold=(r_index == 0))


def insert_table_after(doc, paragraph, headers, rows, widths=None,
                       font_size=10.5):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, value in enumerate(headers):
        table.cell(0, i).text = str(value)
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = str(value)
    format_table(table, font_size=font_size, widths=widths)
    paragraph._p.addnext(table._tbl)
    return table


def add_disclosure(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run('AI Use Disclosure')
    set_run_font(r, size=12, bold=True, color=BLUE)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run('I used ChatGPT to help check the calculations, organize the step-by-step explanations, and format the completed document. I reviewed the final answers before submission.')
    set_run_font(r2, size=10.5)


def create_circuit_diagram():
    width, height = 1900, 980
    image = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(image)
    font_dir = Path('/System/Library/Fonts/Supplemental')
    normal = ImageFont.truetype(str(font_dir / 'Arial.ttf'), 34)
    bold = ImageFont.truetype(str(font_dir / 'Arial Bold.ttf'), 38)
    title = ImageFont.truetype(str(font_dir / 'Arial Bold.ttf'), 46)

    draw.text((80, 35), "F = A'B'C' + AB + AC + BC", font=title, fill='#173B57')
    draw.text((80, 110), 'Input signals and inversions', font=bold, fill='#294E68')
    signals = [('A', "A'"), ('B', "B'"), ('C', "C'")]
    for i, (signal, inverted) in enumerate(signals):
        y = 205 + i * 150
        draw.text((70, y + 20), signal, font=bold, fill='#111111')
        draw.line((120, y + 45, 250, y + 45), fill='#657786', width=5)
        draw.rounded_rectangle((250, y, 450, y + 90), radius=18,
                               outline='#356789', fill='#EAF2F8', width=5)
        label = f'NOT {signal}'
        box = draw.textbbox((0, 0), label, font=normal)
        draw.text((350 - (box[2] - box[0]) / 2, y + 24), label,
                  font=normal, fill='#173B57')
        draw.line((450, y + 45, 550, y + 45), fill='#657786', width=5)
        draw.text((565, y + 20), inverted, font=bold, fill='#247044')

    draw.text((750, 110), 'Product terms', font=bold, fill='#294E68')
    terms = ["A' · B' · C'", 'A · B', 'A · C', 'B · C']
    gate_centers = []
    for i, term in enumerate(terms):
        y = 180 + i * 145
        draw.rounded_rectangle((700, y, 1120, y + 100), radius=24,
                               outline='#356789', fill='#DDEBF7', width=5)
        draw.text((730, y + 18), 'AND', font=bold, fill='#173B57')
        draw.text((850, y + 24), term, font=normal, fill='#111111')
        gate_centers.append(y + 50)

    draw.text((1390, 110), 'Final output', font=bold, fill='#294E68')
    draw.rounded_rectangle((1370, 350, 1640, 500), radius=28,
                           outline='#247044', fill='#E2F0D9', width=6)
    draw.text((1455, 392), 'OR', font=bold, fill='#174D2D')
    for y in gate_centers:
        draw.line((1120, y, 1280, y), fill='#657786', width=5)
        draw.line((1280, y, 1280, 425), fill='#657786', width=5)
        draw.line((1280, 425, 1370, 425), fill='#657786', width=5)
    draw.line((1640, 425, 1770, 425), fill='#247044', width=7)
    draw.text((1790, 400), 'F', font=bold, fill='#174D2D')
    draw.rounded_rectangle((90, 800, 1810, 930), radius=20,
                           outline='#C7D3DC', fill='#F4F7F9', width=3)
    note = "Use NOT gates to make A', B', and C'. The four AND gates make the four terms. The OR gate combines them into F."
    draw.text((145, 842), note, font=normal, fill='#415564')
    image.save(CIRCUIT_PATH, dpi=(300, 300))


def build_boolean_document():
    doc = Document(BOOL_SOURCE)
    p = doc.paragraphs

    answers = {
        5: "Answer:\n(A'BC')' + A(B' + C)\n= A + B' + C + AB' + AC   (De Morgan's law and distribution)\n= A + B' + C   (absorption: A + AX = A)",
        6: "Answer:\nA'B + AB' is the XOR function, written A ⊕ B. It is already in its simplest sum-of-products form because the two terms cover different input cases.",
        7: "Answer:\n(A + B)(A' + C)\n= AA' + AC + A'B + BC\n= AC + A'B + BC\n= AC + A'B   (consensus theorem removes BC)",
        8: "Answer:\n(A + B)(A + C)(B + C)\n= (A + BC)(B + C)\n= AB + AC + BC",
        9: "Answer:\n(A + B)(A' + C)\n= AA' + AC + A'B + BC\n= 0 + AC + A'B + BC\n= AC + A'B.\nThe term BC is redundant by the consensus theorem, so the two sides are equivalent.",
        10: "Answer (five NAND gates):\n1. B' = NAND(B, B)\n2. A' = NAND(A, A)\n3. X = NAND(A, B') = (AB')'\n4. Y = NAND(A', C) = (A'C)'\n5. F = NAND(X, Y) = AB' + A'C",
        19: "Answer:\nF = Σ(0, 3, 5, 6, 7)\n= A'B'C' + BC + AC + AB.\nMinterm 0 is isolated, while 3 pairs with 7 to make BC, 5 pairs with 7 to make AC, and 6 pairs with 7 to make AB.",
        29: "Answer: The simplified expression is F = A'B'C' + AB + AC + BC. The circuit below uses three NOT gates, four AND gates, and one OR gate.",
        31: "Answer: The truth table is shown below.",
        32: "Answer: A minterm listed in Σ receives output 1; every other row receives 0.",
        34: "Answer: AND is 1 only when both inputs are 1.",
        35: "Answer: OR is 1 when at least one input is 1.",
        36: "Answer: NOT reverses A. B does not affect the result.",
        37: "Answer: This is XOR, so the output is 1 when A and B are different.",
        38: "Answer: The output is 1 for minterms 3, 4, and 7.",
    }
    for index, answer in answers.items():
        replace_placeholder(p[index], answer)

    # Start the truth-table section on a clean page so its first table does not
    # split after a single data row.
    p[30].paragraph_format.page_break_before = True

    # Circuit diagram under question 5.
    picture_p = doc.add_paragraph()
    picture_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_p.paragraph_format.space_after = Pt(8)
    picture_p.add_run().add_picture(str(CIRCUIT_PATH), width=Inches(6.25))
    p[29]._p.addnext(picture_p._p)

    rows3 = list(product([0, 1], repeat=3))
    table6a = [(a, b, c, int((a or (not b)) and ((not a) or c)))
               for a, b, c in rows3]
    table6b = [(a, b, c, int((4*a + 2*b + c) in {0, 1, 2, 4, 7}))
               for a, b, c in rows3]
    insert_table_after(doc, p[31], ['A', 'B', 'C', 'F'], table6a,
                       widths=[900, 900, 900, 900])
    insert_table_after(doc, p[32], ['A', 'B', 'C', 'F'], table6b,
                       widths=[900, 900, 900, 900])

    rows2 = list(product([0, 1], repeat=2))
    tables = {
        34: [(a, b, a & b) for a, b in rows2],
        35: [(a, b, a | b) for a, b in rows2],
        36: [(a, b, 1-a) for a, b in rows2],
        37: [(a, b, a ^ b) for a, b in rows2],
        38: [(a, b, c, int((a, b, c) in {(0, 1, 1), (1, 0, 0), (1, 1, 1)}))
             for a, b, c in rows3],
    }
    for index in [34, 35, 36, 37]:
        insert_table_after(doc, p[index], ['A', 'B', 'F'], tables[index],
                           widths=[900, 900, 900])
    insert_table_after(doc, p[38], ['A', 'B', 'C', 'F'], tables[38],
                       widths=[900, 900, 900, 900])
    # Keep the NOT table from splitting after three of its four rows.
    p[36].paragraph_format.page_break_before = True
    add_disclosure(doc)
    doc.save(BOOL_OUT)


def build_adjacency_document():
    doc = Document(ADJ_SOURCE)
    p = doc.paragraphs
    answers = {
        4: "Determine the adjacency matrices of r1 and r2 for the information provided:\nLet A1 = {3, 8, 4, 6}, A2 = {6, 8, 7}, and A3 = {1, 2, 3}. Let R1 be the relation from A1 into A2 defined by r1 = {(x, y)| y − x =2}, and let r2 be the relation from A2 into A3 defined by r2 = {(x, y) | y − x = 1}.\nAnswer: r1 = {(4, 6), (6, 8)}. No value in A3 is one more than a value in A2, so r2 is empty. Rows are source-set values and columns are destination-set values. The matrices are shown below.",
        5: "Write pseudocode to find the shortest path between two vertices in a weighted graph using an adjacency matrix.\nAnswer (Dijkstra's algorithm):\nSET distance for every vertex to infinity\nSET previous for every vertex to null\nSET visited for every vertex to false\nSET distance[start] to 0\nWHILE an unvisited reachable vertex remains\n    SET current to the unvisited vertex with the smallest distance\n    IF current equals target, STOP the loop\n    SET visited[current] to true\n    FOR each vertex neighbor\n        IF matrix[current][neighbor] is an edge AND neighbor is unvisited\n            SET newDistance to distance[current] + matrix[current][neighbor]\n            IF newDistance is less than distance[neighbor]\n                SET distance[neighbor] to newDistance\n                SET previous[neighbor] to current\n            END IF\n        END IF\n    END FOR\nEND WHILE\nFOLLOW previous from target back to start to build the path\nRETURN the path and distance[target]",
        6: "Analyze the time and space complexity of different graph algorithms when using adjacency matrix representation.\nAnswer: The matrix itself always uses O(V²) space. Because every row has V possible neighbors, BFS, DFS, Dijkstra, and Prim all take O(V²) time with this representation. Floyd–Warshall uses O(V³) time. The comparison table is shown below.",
        8: "B = [7 8]\n      [9 10]\n      [11 12]\nAnswer: No. Matrix A is 2 × 3, while B is 3 × 2. Addition requires both matrices to have exactly the same number of rows and columns.",
        9: "Explain whether it is possible to multiply the matrices provided.\nA = [1 2 3] \n       [4 5 6] \n\nB = [7 8]\n       [9 10]\n       [11 12]\nAnswer: Yes. A is 2 × 3 and B is 3 × 2. The inside dimensions match (3 = 3), so AB exists and its result is 2 × 2.",
        11: "B = [7 8]\n       [9 10]\n       [11 12]\nAnswer:\nAB = [1(7)+2(9)+3(11)    1(8)+2(10)+3(12)]\n        [4(7)+5(9)+6(11)    4(8)+5(10)+6(12)]\n     = [58   64]\n        [139 154]",
        12: "Use matrix methods to solve the system of linear equations provided.\n3x + 2y − z = 1\n2x − y + 3z = -3\nx + y + z = 2\nAnswer: Write the augmented matrix and use row operations:\n[3  2 -1 |  1]\n[2 -1  3 | -3]\n[1  1  1 |  2]\nSwap the last row to the top, then eliminate x. This gives the two equations 3y − z = 7 and y + 4z = 5. Solving them gives z = 8/13, y = 33/13, and x = -15/13.\nSolution: (x, y, z) = (-15/13, 33/13, 8/13).",
        13: "Explain how matrices can be used to represent graphs. \nAnswer: Give every vertex a row and a column. Put 1 in position (i, j) when an unweighted edge connects vertex i to vertex j, and put 0 when no edge exists. For a weighted graph, store the edge weight instead of 1. An undirected graph has a symmetric matrix because an edge works in both directions.",
        15: "Answer: Matrix operations help computers handle large groups of connected values at once. Adjacency matrices support graph searches and path counting. Matrix multiplication is used in graphics, machine learning, image processing, simulations, and network analysis. Powers of an adjacency matrix can count walks of a specific length between vertices.",
        16: "Determine the determinant of A for the matrix provided.\nA = [2 3 1]\n       [1 4 6]\n       [6 8 3]\nAnswer:\ndet(A) = 2(4·3 − 6·8) − 3(1·3 − 6·6) + 1(1·8 − 4·6)\n= 2(-36) − 3(-33) − 16\n= -72 + 99 - 16\n= 11.",
        17: "Determine the inverse of A for the matrix provided.\nA = [2 3 1]\n       [1 4 6]\n       [6 8 3]\nAnswer: Since det(A) = 11, the inverse exists.\nA⁻¹ = (1/11)[-36  -1  14]\n              [ 33   0  -6]\n              [-16   2   5]",
        21: "Answer: Use det(A − λI) = 0:\n|(3−λ)  1|\n|  1   (3−λ)| = (3−λ)² − 1 = 0.\nThe eigenvalues are λ = 4 and λ = 2.\nFor λ = 4, an eigenvector is [1, 1]ᵀ.\nFor λ = 2, an eigenvector is [1, -1]ᵀ.\nAny nonzero multiple of either vector is also valid.",
    }
    for index, answer in answers.items():
        paragraph = p[index]
        paragraph.clear()
        r = paragraph.add_run(answer)
        set_run_font(r, size=12)
        paragraph.paragraph_format.space_after = Pt(6)

    # Relation matrices after question 1.
    r1 = insert_table_after(doc, p[4], ['A1 / A2', '6', '8', '7'],
                            [('3', 0, 0, 0), ('8', 0, 0, 0),
                             ('4', 1, 0, 0), ('6', 0, 1, 0)],
                            widths=[1500, 900, 900, 900])
    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rl = label.add_run('Matrix for r1')
    set_run_font(rl, size=10.5, bold=True, color=BLUE)
    p[4]._p.addnext(label._p)
    label._p.addnext(r1._tbl)
    r2 = doc.add_table(rows=1, cols=4)
    for i, value in enumerate(['A2 / A3', '1', '2', '3']):
        r2.cell(0, i).text = value
    for value in ['6', '8', '7']:
        cells = r2.add_row().cells
        cells[0].text = value
        for i in range(1, 4):
            cells[i].text = '0'
    format_table(r2, widths=[1500, 900, 900, 900])
    label2 = doc.add_paragraph()
    rl2 = label2.add_run('Matrix for r2')
    set_run_font(rl2, size=10.5, bold=True, color=BLUE)
    r1._tbl.addnext(label2._p)
    label2._p.addnext(r2._tbl)

    complexity_rows = [
        ('Check one edge', 'O(1)', 'O(1)'),
        ('List neighbors of one vertex', 'O(V)', 'O(V) output at most'),
        ('BFS or DFS', 'O(V²)', 'O(V) extra'),
        ('Dijkstra (linear selection)', 'O(V²)', 'O(V) extra'),
        ('Prim (linear selection)', 'O(V²)', 'O(V) extra'),
        ('Floyd–Warshall', 'O(V³)', 'O(V²) distance matrix'),
    ]
    insert_table_after(doc, p[6], ['Operation/Algorithm', 'Time', 'Extra Space'],
                       complexity_rows, widths=[4200, 1500, 2500], font_size=9.5)
    # Keep the required AI acknowledgment with the final answer instead of
    # allowing a short disclosure to become a page by itself.
    disclosure = p[21].add_run(' AI use: ChatGPT helped check the calculations and formatting; I reviewed the final answers.')
    set_run_font(disclosure, size=7.5, bold=False, color=BLUE)
    doc.save(ADJ_OUT)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement('w:tblHeader')
    header.set(qn('w:val'), 'true')
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    cant_split.set(qn('w:val'), 'true')
    tr_pr.append(cant_split)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, name='Calibri', size=9, color='666666')
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, end])


def build_explainer_guide():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)
    header = section.header.paragraphs[0]
    header.text = ''
    hr = header.add_run('CS 505  |  Module Seven Simple Explainer')
    set_run_font(hr, name='Calibri', size=9, bold=True, color='666666')
    footer = section.footer.paragraphs[0]
    footer.text = ''
    add_page_number(footer)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    tr = title.add_run('Boolean Algebra and Matrices: Simple Explainer Tables')
    set_run_font(tr, name='Calibri', size=19, bold=True, color=BLUE)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(6)
    sr = sub.add_run('Plain-language review written at about a 10th-grade level')
    set_run_font(sr, name='Calibri', size=10.5, italic=True, color='666666')

    rows = [
        ('Boolean: prime mark', "A' means NOT A.", 'Flip the value: 0 becomes 1, and 1 becomes 0.', "If A = 1, then A' = 0."),
        ('Boolean: AND', 'AND is 1 only when every required input is 1.', 'Think “all conditions must be true.”', 'AB = 1 only for A = 1 and B = 1.'),
        ('Boolean: OR', 'OR is 1 when at least one input is 1.', 'Think “one or more is enough.”', 'A + B = 0 only when both are 0.'),
        ('De Morgan’s law', 'When a whole group is NOT, flip AND/OR and flip every input.', "(A'BC')' becomes A + B' + C.", 'NOT of a product becomes a sum of NOTs.'),
        ('Absorption law', 'A bigger term can be swallowed by a simpler term.', 'Use A + AX = A.', "A + AB' becomes A."),
        ('XOR', 'XOR is 1 when the two inputs are different.', "Recognize A'B + AB'.", "A'B + AB' = A ⊕ B."),
        ('Consensus theorem', 'A third overlapping term may be unnecessary.', "XY + X'Z + YZ = XY + X'Z.", "AC + A'B + BC becomes AC + A'B."),
        ('NAND-only circuit', 'A NAND gate can build NOT, AND, and OR behavior.', 'Tie inputs together for NOT, then use a final NAND to combine complemented products.', "F = NAND((AB')', (A'C)')."),
        ('Truth table', 'A truth table tries every possible input combination.', 'Three inputs create 2³ = 8 rows.', 'List 000 through 111 and calculate F.'),
        ('Minterm notation', 'Σ lists the row numbers where F = 1.', 'Convert ABC to binary row numbers.', 'Σ(0,1,2,4,7) means those five rows output 1.'),
        ('Adjacency matrix', 'Rows and columns represent vertices or set values.', 'Put 1 for a connection and 0 for no connection.', 'The cell in row i, column j describes i → j.'),
        ('Relation matrix', 'A rule decides whether each ordered pair belongs.', 'Test each row value x with each column value y.', 'For y − x = 2, (4,6) belongs because 6 − 4 = 2.'),
        ('Shortest path', 'Dijkstra repeatedly chooses the closest unfinished vertex.', 'Track distance, previous vertex, and visited status.', 'An adjacency matrix version takes O(V²) time.'),
        ('Big O with a matrix', 'The graph matrix stores every possible pair of vertices.', 'A V by V matrix has V² cells.', 'Matrix storage is O(V²), even for a sparse graph.'),
        ('Matrix addition', 'Matrices add only when their shapes match.', 'Match row count and column count.', 'A 2×3 matrix cannot add to a 3×2 matrix.'),
        ('Matrix multiplication', 'The inside dimensions must match.', 'Multiply each row of A by each column of B.', '(2×3)(3×2) works and produces 2×2.'),
        ('Linear system', 'A matrix stores the coefficients and answers together.', 'Use row operations to isolate x, y, and z.', 'Here x = -15/13, y = 33/13, z = 8/13.'),
        ('Determinant', 'The determinant is one number that tells whether a square matrix is invertible.', 'Expand by a row or use row reduction.', 'det(A) = 11, so the inverse exists.'),
        ('Inverse', 'The inverse undoes the original matrix.', 'A matrix has an inverse only when its determinant is not 0.', 'A⁻¹A = I, the identity matrix.'),
        ('Eigenvalue/eigenvector', 'An eigenvector keeps its direction after multiplication; the eigenvalue tells the scale.', 'Solve det(A − λI) = 0, then solve for a matching vector.', 'For [[3,1],[1,3]], λ = 4 uses [1,1] and λ = 2 uses [1,-1].'),
    ]
    table = doc.add_table(rows=1, cols=4)
    headers = ['Topic', 'Simple meaning', 'What you do', 'Example or key answer']
    for i, text in enumerate(headers):
        table.cell(0, i).text = text
    for row_data in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row_data):
            cells[i].text = text
    widths = [1900, 3200, 3200, 3700]
    format_table(table, font_size=8.25, header_fill=BLUE, widths=widths, indent=120)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        prevent_row_split(row)
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, name='Calibri', size=9, bold=True, color='FFFFFF')
    for r_index, row in enumerate(table.rows[1:], start=1):
        if r_index % 2 == 0:
            for cell in row.cells:
                shade_cell(cell, LIGHT_GRAY)
        for c_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_run_font(run, name='Calibri', size=8.25,
                                 bold=(c_index == 0),
                                 color=(BLUE if c_index == 0 else '000000'))

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(2)
    nr = note.add_run('AI Use Disclosure: ')
    set_run_font(nr, name='Calibri', size=9, bold=True, color=BLUE)
    nr2 = note.add_run('I used ChatGPT to help verify the calculations and turn the main rules into a simple study table. I reviewed the final examples and answers.')
    set_run_font(nr2, name='Calibri', size=9)
    doc.save(GUIDE_OUT)


def main():
    ROOT.mkdir(parents=True, exist_ok=True)
    create_circuit_diagram()
    build_boolean_document()
    build_adjacency_document()
    build_explainer_guide()
    print(BOOL_OUT)
    print(ADJ_OUT)
    print(GUIDE_OUT)


if __name__ == '__main__':
    main()
